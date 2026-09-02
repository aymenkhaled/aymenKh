from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\portfolio\output\cv\Aymen_Khaled_CV.docx")

NAVY = RGBColor(11, 37, 69)
TEAL = RGBColor(15, 139, 141)
INK = RGBColor(26, 36, 48)
MUTED = RGBColor(82, 96, 109)
LIGHT = RGBColor(218, 225, 232)
WHITE = RGBColor(255, 255, 255)


def set_font(run, size=10, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_spacing(paragraph, before=0, after=0, line=1.0, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next


def add_bottom_rule(paragraph, color="0F8B8D", size="12", space="6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_hyperlink(paragraph, text, url, color=TEAL):
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run_element = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color_element = OxmlElement("w:color")
    color_element.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    r_pr.append(color_element)
    font = OxmlElement("w:rFonts")
    font.set(qn("w:ascii"), "Calibri")
    font.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(font)
    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "18")
    r_pr.append(size)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run_element.extend([r_pr, text_element])
    hyperlink.append(run_element)
    paragraph._p.append(hyperlink)


def configure_bullets(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "40")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, num_id, text):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_element = OxmlElement("w:numId")
    num_id_element.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_element])
    p_pr.append(num_pr)
    run = p.add_run(text)
    set_font(run, size=9.4)
    return p


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=7, after=4, keep_next=True)
    r = p.add_run(text.upper())
    set_font(r, size=11, bold=True, color=NAVY)
    add_bottom_rule(p, color="D9E1E8", size="5", space="2")
    return p


def add_role(doc, title, company, location, dates, bullets, num_id):
    p = doc.add_paragraph()
    set_spacing(p, before=3, after=1, keep_next=True)
    r = p.add_run(f"{title} | {company}")
    set_font(r, size=10.2, bold=True, color=NAVY)
    r = p.add_run(f"  •  {location}")
    set_font(r, size=9.3, color=MUTED)
    p2 = doc.add_paragraph()
    set_spacing(p2, after=2, keep_next=True)
    r = p2.add_run(dates)
    set_font(r, size=9.2, bold=True, color=TEAL)
    for item in bullets:
        add_bullet(doc, num_id, item)


def add_project(doc, title, stack, text):
    p = doc.add_paragraph()
    set_spacing(p, before=3, after=1, keep_next=True)
    r = p.add_run(title)
    set_font(r, size=10, bold=True, color=NAVY)
    r = p.add_run(f"  |  {stack}")
    set_font(r, size=9.1, color=TEAL)
    p2 = doc.add_paragraph()
    set_spacing(p2, after=3)
    r = p2.add_run(text)
    set_font(r, size=9.4)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.45)
section.bottom_margin = Inches(0.45)
section.left_margin = Inches(0.68)
section.right_margin = Inches(0.68)
section.header_distance = Inches(0.2)
section.footer_distance = Inches(0.2)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(9.4)
normal.font.color.rgb = INK
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.0

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_spacing(header, after=0)
hr = header.add_run("AYMEN KHALED  •  FULL-STACK DEVELOPER")
set_font(hr, size=8, bold=True, color=MUTED)

footer = section.footer.paragraphs[0]
add_page_number(footer)

num_id = configure_bullets(doc)

name = doc.add_paragraph()
set_spacing(name, after=1)
r = name.add_run("AYMEN KHALED")
set_font(r, size=25, bold=True, color=NAVY)

headline = doc.add_paragraph()
set_spacing(headline, after=4)
r = headline.add_run("FULL-STACK DEVELOPER")
set_font(r, size=12.5, bold=True, color=TEAL)
r = headline.add_run("  |  SaaS • Web • Mobile • Windows Desktop • AI & Automation")
set_font(r, size=9.4, color=MUTED)

contact = doc.add_paragraph()
set_spacing(contact, after=7)
r = contact.add_run("Sfax, Tunisia  |  +216 26 286 045  |  khaledaymen850@gmail.com\n")
set_font(r, size=9.1, color=INK)
add_hyperlink(contact, "aymen-kh.vercel.app", "https://aymen-kh.vercel.app/")
r = contact.add_run("  |  ")
set_font(r, size=9.1, color=MUTED)
add_hyperlink(contact, "linkedin.com/in/aymen-khaled-652724236", "https://www.linkedin.com/in/aymen-khaled-652724236/")
r = contact.add_run("  |  ")
set_font(r, size=9.1, color=MUTED)
add_hyperlink(contact, "github.com/aymenkhaled", "https://github.com/aymenkhaled")
add_bottom_rule(contact, color="0F8B8D", size="10", space="5")

add_section_heading(doc, "Professional Summary")
p = doc.add_paragraph()
set_spacing(p, after=3)
r = p.add_run(
    "Full-stack developer building SaaS, web, mobile, and Windows desktop products. "
    "Experienced with multi-tenant platforms, real-time features, AI integrations, billing, automation, and offline-first applications, "
    "with a strong focus on React, TypeScript, Node.js, FastAPI, and PostgreSQL."
)
set_font(r, size=9.6)

add_section_heading(doc, "Technical Skills")
skills = [
    ("Languages", "TypeScript, JavaScript, Python, Rust"),
    ("Frontend", "React, Next.js, React Native, Redux, Tailwind CSS"),
    ("Backend & APIs", "Node.js, Express, FastAPI, Django, REST APIs, WebSockets, JWT, OAuth"),
    ("Data", "PostgreSQL, MongoDB, SQLite, SQL Server, Redis, Firebase"),
    ("AI & Automation", "OpenAI, Groq, LangChain, Whisper, web scraping, enrichment workflows"),
    ("Delivery & Integrations", "Tauri, Docker, Git, GitHub Actions, Vercel, AWS, Stripe, PayPal, Slack, Asana, Basecamp"),
]
for label, value in skills:
    p = doc.add_paragraph()
    set_spacing(p, after=1)
    r = p.add_run(f"{label}: ")
    set_font(r, size=9.3, bold=True, color=NAVY)
    r = p.add_run(value)
    set_font(r, size=9.3)

add_section_heading(doc, "Professional Experience")
add_role(
    doc,
    "Full-Stack Developer",
    "Everything to Gain",
    "Remote",
    "Aug 2025 – Present",
    [
        "Contribute to Strategy Navigator, JourneyAI, and SaleSide AI across frontend, backend, AI, and product integrations.",
        "Integrated 65+ AI tools and delivered billing workflows with Stripe and PayPal.",
        "Built 10+ AI assistants with real-time chat, speech-to-text, summaries, and follow-up workflows.",
    ],
    num_id,
)
add_role(
    doc,
    "Full-Stack Developer Intern",
    "Aziin Engineering Solution",
    "Sfax, Tunisia",
    "Feb 2025 – May 2025",
    [
        "Developed an e-learning platform with MERN features, authentication, AI chatbot support, and WebSocket interactions.",
        "Built course and content workflows and improved application responsiveness for real-time use.",
    ],
    num_id,
)

add_role(
    doc,
    "Mobile Developer Intern",
    "SAC Marquage",
    "Tunisia",
    "Sep 2024 – Oct 2024",
    [
        "Built a React Native RFID application and Django REST APIs for tag workflows and device synchronization.",
    ],
    num_id,
)
add_role(
    doc,
    "Full-Stack Web Developer Intern",
    "Proged",
    "Tunisia",
    "Jul 2022 – Aug 2022",
    [
        "Contributed to an e-commerce system covering product catalog, payments, order workflows, and maintenance.",
    ],
    num_id,
)

add_section_heading(doc, "Selected Projects")
add_project(
    doc,
    "Library Stock — Windows POS & Inventory",
    "React, TypeScript, Tauri, Rust, SQLite",
    "Offline-first Windows application for bookstore sales, stock, purchases, returns, barcode workflows, ticket printing, backup, and restore. Packaged with transactional Rust commands and local SQLite storage.",
)
add_project(
    doc,
    "BERBRY — 3D E-commerce Platform",
    "React, TypeScript, Three.js, Express, PostgreSQL",
    "Responsive fashion storefront and admin system with product variants, cart and checkout, order management, uploads, and interactive 3D product experiences.",
)
add_project(
    doc,
    "AutoTN — Trilingual Automotive Marketplace",
    "React Native, Expo, Express, PostgreSQL",
    "Arabic, French, and English mobile marketplace with AI diagnostics, VIN decoding, price estimation, seller profiles, listings, and WhatsApp contact flows.",
)
add_section_heading(doc, "Education")
for title, school, date in [
    ("BSc Big Data & Analytics", "ISIMS, Sfax, Tunisia", "2022 – 2025"),
    ("Software Engineering Diploma", "ESPS, Tunisia", "2022"),
]:
    p = doc.add_paragraph()
    set_spacing(p, after=2)
    r = p.add_run(title)
    set_font(r, size=9.7, bold=True, color=NAVY)
    r = p.add_run(f"  |  {school}  |  {date}")
    set_font(r, size=9.3, color=MUTED)

add_section_heading(doc, "Certifications & Languages")
p = doc.add_paragraph()
set_spacing(p, after=2)
r = p.add_run("Certifications: ")
set_font(r, size=9.3, bold=True, color=NAVY)
r = p.add_run("JavaScript Algorithms — freeCodeCamp; Responsive Web Design — freeCodeCamp; Young Entrepreneur Rookies")
set_font(r, size=9.3)
p = doc.add_paragraph()
set_spacing(p, after=2)
r = p.add_run("Languages: ")
set_font(r, size=9.3, bold=True, color=NAVY)
r = p.add_run("Arabic — Native; English — Advanced; French — Intermediate (B1)")
set_font(r, size=9.3)

for paragraph in doc.paragraphs:
    paragraph.paragraph_format.widow_control = True

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
