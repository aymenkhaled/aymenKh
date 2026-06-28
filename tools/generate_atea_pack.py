from __future__ import annotations

import io
import re
import shutil
import zipfile
from datetime import date
from pathlib import Path

from PIL import Image as PILImage
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.lib.utils import ImageReader


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
OUTPUT_ROOT = WORKSPACE / "output"
PDF_DIR = OUTPUT_ROOT / "pdf"
DOCX_DIR = OUTPUT_ROOT / "docx"
CHECKLIST_DIR = OUTPUT_ROOT / "checklists"
SCREENSHOT_DIR = (
    WORKSPACE
    / "tmp"
    / "evidence_analysis"
    / "lms_video_screenshots_and_analysis"
    / "screenshots_by_section"
)
FINAL_ZIP = OUTPUT_ROOT / "ATEA_Aymen_Khaled_Submission_Pack.zip"

NAVY = colors.HexColor("#0B2E4F")
BLUE = colors.HexColor("#1769AA")
CYAN = colors.HexColor("#0E7C86")
LIGHT_BLUE = colors.HexColor("#EAF3F8")
LIGHT_GREY = colors.HexColor("#F2F4F6")
MID_GREY = colors.HexColor("#667085")
GREEN = colors.HexColor("#16794A")
AMBER = colors.HexColor("#A15C00")
RED = colors.HexColor("#A12622")


def required_artifacts() -> list[str]:
    return [
        "pdf/CV_Aymen_Khaled_ATEA_FR.pdf",
        "docx/CV_Aymen_Khaled_ATEA_FR.docx",
        "pdf/Dossier_Expert_Technique_Aymen_Khaled_ATEA.pdf",
        "pdf/Annexe_Visuelle_Demonstrateur_LMS_ATEA.pdf",
        "pdf/Architecture_LMS_Sur_Mesure_ATEA.pdf",
        "docx/Corrections_Proposition_Mehdi_ATEA.docx",
        "pdf/Corrections_Proposition_Mehdi_ATEA.pdf",
        "checklists/Pieces_Justificatives_Manquantes_ATEA.md",
        "README_ENVOI.md",
    ]


def proposed_role() -> str:
    return "Développeur Full-Stack - Responsable du développement technique de la plateforme LMS"


PROFILE = {
    "name": "Aymen Khaled",
    "role": proposed_role(),
    "location": "Mahdia, Tunisie",
    "phone": "+216 26 286 045",
    "email": "khaledaymen850@gmail.com",
    "portfolio": "https://aymen-kh.vercel.app",
    "linkedin": "https://linkedin.com/in/aymen-khaled-652724236",
    "github": "https://github.com/aymenkhaled",
}


EDUCATION = [
    ("Licence en Big Data et Analyse de Données", "ISIMS, Sfax", "2022 - 2025"),
    ("Diplôme en Génie Logiciel", "ESPS, Tunisie", "2022"),
]

CERTIFICATIONS = [
    "JavaScript Algorithms - freeCodeCamp",
    "Responsive Web Design - freeCodeCamp",
    "Young Entrepreneur Rookies",
]

EXPERIENCE = [
    (
        "Développeur Web Full-Stack",
        "Everything to Gain - à distance",
        "Février 2025 - présent",
        [
            "Développement de plateformes SaaS, d'intégrations d'IA et de flux multi-utilisateurs.",
            "Contribution à Strategy Navigator, JourneyAI et SaleSide AI selon le CV fourni.",
            "Intégrations API, authentification par rôles, temps réel, paiements et automatisation.",
        ],
    ),
    (
        "Développeur Full-Stack - stage",
        "Aziin Engineering Solution",
        "2024",
        [
            "Développement d'un démonstrateur LMS avec React, Node.js, MongoDB et WebSockets.",
            "Fonctions présentées: cours, leçons, sessions en direct, tableau de bord, quiz assisté par IA et administration.",
            "Le statut retenu dans ce dossier est démonstrateur de stage, faute d'attestation de réception client fournie.",
        ],
    ),
    (
        "Développeur Mobile - stage",
        "SAC Marquage",
        "Septembre 2024 - octobre 2024",
        [
            "Application React Native connectée à des dispositifs RFID par Bluetooth et NFC.",
            "APIs Django REST pour la gestion des tags et la synchronisation des équipements.",
        ],
    ),
    (
        "Développeur Web Full-Stack - stage",
        "Proged",
        "Juillet 2022 - août 2022",
        [
            "Application e-commerce avec backend .NET, frontend React, SQL Server et MongoDB.",
            "Contribution aux flux catalogue, commande et paiement.",
        ],
    ),
]

PROJECTS = [
    {
        "title": "EduNova - Démonstrateur LMS sur mesure",
        "status": "Référence directement similaire - démonstrateur de stage",
        "period": "2024",
        "role": "Développeur Full-Stack",
        "stack": "React, Node.js, MongoDB, WebSockets, intégrations IA",
        "scope": (
            "Gestion des cours et leçons, catalogue, sessions en direct, communication temps réel, "
            "tableaux de bord apprenant et administrateur, parcours assistés par IA et outils d'évaluation."
        ),
        "evidence": "Vidéo de démonstration, 25 captures d'écran et URL publique du démonstrateur.",
        "url": "https://e-learning-five-tau.vercel.app",
        "limits": (
            "Les pièces fournies ne comprennent pas d'attestation client, de procès-verbal de réception, "
            "de rapport de tests de sécurité ni de preuve de charge."
        ),
    },
    {
        "title": "Strategy Navigator - Plateforme SaaS multi-tenant",
        "status": "Référence technique complémentaire - produit public",
        "period": "Depuis 2025 selon le CV",
        "role": "Développeur Web Full-Stack",
        "stack": "React, Node.js, API, RBAC, intégrations IA",
        "scope": "Architecture multi-tenant, contrôle d'accès par rôles, tableaux de bord et intégrations externes.",
        "evidence": "Produit public et description dans le CV/portfolio.",
        "url": "https://strategynavigator.ai",
        "limits": "Joindre une attestation employeur pour transformer cette référence en preuve nominative forte.",
    },
    {
        "title": "E-Commerce Dashboard - Administration et analytique",
        "status": "Référence technique complémentaire - code public",
        "period": "Projet portfolio",
        "role": "Développeur Full-Stack",
        "stack": "React, Node.js, PostgreSQL",
        "scope": "Tableaux de bord, suivi d'activité, inventaire, commandes et vues analytiques.",
        "evidence": "Démonstration publique et dépôt GitHub public au nom d'Aymen Khaled.",
        "url": "https://dashboard-rho-gray-61.vercel.app/dashboard\nhttps://github.com/aymenkhaled/dashboard",
        "limits": "Référence utile pour l'administration et le reporting, mais non équivalente à une mission LMS.",
    },
]


def register_fonts() -> None:
    candidates = {
        "AteaSans": Path(r"C:\Windows\Fonts\arial.ttf"),
        "AteaSans-Bold": Path(r"C:\Windows\Fonts\arialbd.ttf"),
        "AteaSans-Italic": Path(r"C:\Windows\Fonts\ariali.ttf"),
    }
    for name, path in candidates.items():
        if path.exists() and name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(name, str(path)))


def styles():
    register_fonts()
    base = getSampleStyleSheet()
    body = ParagraphStyle(
        "AteaBody",
        parent=base["BodyText"],
        fontName="AteaSans",
        fontSize=9.2,
        leading=13,
        textColor=colors.HexColor("#1D2939"),
        alignment=TA_JUSTIFY,
        spaceAfter=5,
    )
    return {
        "body": body,
        "small": ParagraphStyle("AteaSmall", parent=body, fontSize=7.6, leading=10, textColor=MID_GREY),
        "caption": ParagraphStyle("AteaCaption", parent=body, fontSize=8, leading=10.5, textColor=NAVY),
        "h1": ParagraphStyle(
            "AteaH1",
            parent=base["Heading1"],
            fontName="AteaSans-Bold",
            fontSize=22,
            leading=26,
            textColor=NAVY,
            spaceAfter=12,
        ),
        "h2": ParagraphStyle(
            "AteaH2",
            parent=base["Heading2"],
            fontName="AteaSans-Bold",
            fontSize=14,
            leading=17,
            textColor=BLUE,
            spaceBefore=8,
            spaceAfter=7,
        ),
        "h3": ParagraphStyle(
            "AteaH3",
            parent=base["Heading3"],
            fontName="AteaSans-Bold",
            fontSize=10.5,
            leading=13,
            textColor=NAVY,
            spaceBefore=5,
            spaceAfter=3,
        ),
        "title": ParagraphStyle(
            "AteaTitle",
            parent=base["Title"],
            fontName="AteaSans-Bold",
            fontSize=28,
            leading=33,
            textColor=NAVY,
            alignment=TA_LEFT,
        ),
        "subtitle": ParagraphStyle(
            "AteaSubtitle",
            parent=body,
            fontName="AteaSans",
            fontSize=12,
            leading=16,
            textColor=CYAN,
            alignment=TA_LEFT,
        ),
        "center": ParagraphStyle("AteaCenter", parent=body, alignment=TA_CENTER),
        "bullet": ParagraphStyle(
            "AteaBullet",
            parent=body,
            leftIndent=12,
            firstLineIndent=-8,
            bulletIndent=0,
            spaceAfter=3,
        ),
        "callout": ParagraphStyle(
            "AteaCallout",
            parent=body,
            fontName="AteaSans-Bold",
            textColor=NAVY,
            backColor=LIGHT_BLUE,
            borderColor=colors.HexColor("#B8D7E8"),
            borderWidth=0.6,
            borderPadding=8,
            leading=13,
            spaceBefore=6,
            spaceAfter=8,
        ),
    }


def p(text: str, style, **kwargs) -> Paragraph:
    safe = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(safe, style, **kwargs)


def rich(text: str, style) -> Paragraph:
    return Paragraph(text, style)


def bullets(items: list[str], st) -> list[Paragraph]:
    return [p(f"- {item}", st["bullet"]) for item in items]


def footer(canvas_obj, doc) -> None:
    canvas_obj.saveState()
    width, _ = A4
    canvas_obj.setStrokeColor(colors.HexColor("#D0D5DD"))
    canvas_obj.line(1.7 * cm, 1.25 * cm, width - 1.7 * cm, 1.25 * cm)
    canvas_obj.setFont("AteaSans", 7.5)
    canvas_obj.setFillColor(MID_GREY)
    canvas_obj.drawString(1.7 * cm, 0.85 * cm, "Dossier ATEA - Aymen Khaled - 28/06/2026")
    canvas_obj.drawRightString(width - 1.7 * cm, 0.85 * cm, f"Page {doc.page}")
    canvas_obj.restoreState()


def make_doc(path: Path, title: str, author: str = "Aymen Khaled") -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)
    doc.core_properties.title = title
    doc.core_properties.author = author
    doc.core_properties.subject = "Dossier technique ATEA"
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    for style_name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(11, 46, 79)
    path.parent.mkdir(parents=True, exist_ok=True)
    return doc


def docx_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(4)


def docx_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)
    paragraph.paragraph_format.space_after = Pt(2)


def build_cv_docx(path: Path) -> None:
    doc = make_doc(path, "CV ATEA - Aymen Khaled")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(PROFILE["name"].upper())
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 46, 79)
    role = doc.add_paragraph()
    role.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = role.add_run(PROFILE["role"])
    r.bold = True
    r.font.color.rgb = RGBColor(14, 124, 134)
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.add_run(
        f"{PROFILE['location']} | {PROFILE['phone']} | {PROFILE['email']}\n"
        f"{PROFILE['portfolio']} | {PROFILE['linkedin']} | {PROFILE['github']}"
    )

    docx_heading(doc, "PROFIL PROFESSIONNEL", 1)
    doc.add_paragraph(
        "Développeur Full-Stack orienté plateformes SaaS, intégrations d'IA et applications web/mobile. "
        "Dispose du code source d'un démonstrateur LMS similaire au besoin ATEA et propose son audit, "
        "son adaptation et son industrialisation. Son rôle porte sur le développement technique; "
        "la scénarisation pédagogique demeure sous la responsabilité des experts pédagogiques."
    )

    docx_heading(doc, "COMPÉTENCES PERTINENTES POUR LA MISSION", 1)
    for item in [
        "Frontend: React, Next.js, TypeScript, Redux, Tailwind CSS, interfaces responsives.",
        "Backend: Node.js, Express, FastAPI, Django, APIs REST, GraphQL et WebSockets.",
        "Données: PostgreSQL, MongoDB, SQL Server, Redis et Supabase.",
        "Sécurité applicative: authentification JWT/OAuth2, contrôle d'accès par rôles et validation des entrées.",
        "DevOps: Docker, Jenkins, GitHub Actions, Linux, Nginx, CI/CD et notions AWS S3/EC2.",
        "LMS: cours, leçons, contenus, sessions en direct, progression, évaluations, tableaux de bord et administration.",
    ]:
        docx_bullet(doc, item)

    docx_heading(doc, "EXPÉRIENCE PROFESSIONNELLE", 1)
    for role_name, company, period, achievements in EXPERIENCE:
        heading = doc.add_paragraph()
        hr = heading.add_run(f"{role_name} | {company}")
        hr.bold = True
        heading.add_run(f"\n{period}").italic = True
        for achievement in achievements:
            docx_bullet(doc, achievement)

    doc.add_section(WD_SECTION.NEW_PAGE)
    docx_heading(doc, "RÉFÉRENCES TECHNIQUES", 1)
    for project in PROJECTS:
        h = doc.add_paragraph()
        rr = h.add_run(project["title"])
        rr.bold = True
        rr.font.color.rgb = RGBColor(23, 105, 170)
        doc.add_paragraph(f"Statut: {project['status']} | Période: {project['period']}")
        docx_bullet(doc, f"Rôle: {project['role']}")
        docx_bullet(doc, f"Technologies: {project['stack']}")
        docx_bullet(doc, f"Périmètre: {project['scope']}")
        docx_bullet(doc, f"Preuve disponible: {project['evidence']}")
        docx_bullet(doc, f"Lien: {project['url']}")

    docx_heading(doc, "FORMATION", 1)
    for degree, school, period in EDUCATION:
        docx_bullet(doc, f"{degree} | {school} | {period}")

    docx_heading(doc, "CERTIFICATIONS DÉCLARÉES", 1)
    for certificate in CERTIFICATIONS:
        docx_bullet(doc, certificate)
    note = doc.add_paragraph()
    rn = note.add_run(
        "Important: les copies des diplômes et certificats n'ont pas été fournies avec les fichiers analysés. "
        "Elles doivent être jointes séparément avant la soumission."
    )
    rn.bold = True
    rn.font.color.rgb = RGBColor(161, 92, 0)

    docx_heading(doc, "LANGUES", 1)
    doc.add_paragraph("Arabe: langue maternelle | Anglais: avancé | Français: intermédiaire B1")
    docx_heading(doc, "RÔLE PROPOSÉ DANS LA MISSION ATEA", 1)
    for item in [
        "Audit du code source LMS existant et analyse des écarts par rapport au cahier des charges.",
        "Adaptation des espaces apprenant, formateur, administrateur, référent et expert évaluateur.",
        "Développement des APIs, tableaux de bord, rapports, évaluations, certificats et intégrations.",
        "Tests fonctionnels, multi-rôles, sécurité applicative, performance et documentation technique.",
        "Déploiement, transfert de compétences et support de démarrage.",
    ]:
        docx_bullet(doc, item)
    doc.save(path)


def build_cv_pdf(path: Path) -> None:
    st = styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=1.55 * cm,
        rightMargin=1.55 * cm,
        topMargin=1.45 * cm,
        bottomMargin=1.55 * cm,
        title="CV ATEA - Aymen Khaled",
        author="Aymen Khaled",
    )
    story = [
        rich("<font color='#0B2E4F'><b>AYMEN KHALED</b></font>", ParagraphStyle("cvname", parent=st["title"], fontSize=24, leading=26, alignment=TA_CENTER)),
        rich(f"<font color='#0E7C86'><b>{PROFILE['role']}</b></font>", ParagraphStyle("cvrole", parent=st["subtitle"], alignment=TA_CENTER, fontSize=10.7)),
        p(f"{PROFILE['location']} | {PROFILE['phone']} | {PROFILE['email']}", st["center"]),
        p(f"{PROFILE['portfolio']} | {PROFILE['linkedin']} | {PROFILE['github']}", st["center"]),
        Spacer(1, 5),
        p("PROFIL PROFESSIONNEL", st["h2"]),
        p(
            "Développeur Full-Stack orienté plateformes SaaS, intégrations d'IA et applications web/mobile. "
            "Dispose du code source d'un démonstrateur LMS similaire au besoin ATEA et propose son audit, "
            "son adaptation et son industrialisation. Son rôle porte sur le développement technique; "
            "la scénarisation pédagogique demeure sous la responsabilité des experts pédagogiques.",
            st["body"],
        ),
        p("COMPÉTENCES PERTINENTES", st["h2"]),
    ]
    story.extend(
        bullets(
            [
                "Frontend: React, Next.js, TypeScript, Redux, Tailwind CSS et interfaces responsives.",
                "Backend: Node.js, Express, FastAPI, Django, APIs REST, GraphQL et WebSockets.",
                "Données: PostgreSQL, MongoDB, SQL Server, Redis et Supabase.",
                "Sécurité: JWT/OAuth2, contrôle d'accès par rôles et validation applicative.",
                "DevOps: Docker, Jenkins, GitHub Actions, Linux, Nginx, CI/CD et notions AWS S3/EC2.",
                "LMS: cours, leçons, contenus, sessions en direct, progression, évaluations et administration.",
            ],
            st,
        )
    )
    story.append(p("EXPÉRIENCE PROFESSIONNELLE", st["h2"]))
    for role_name, company, period, achievements in EXPERIENCE:
        block = [
            rich(f"<b>{role_name}</b> | {company}<br/><font color='#667085'><i>{period}</i></font>", st["h3"]),
            *bullets(achievements, st),
        ]
        story.append(KeepTogether(block))
    story.append(PageBreak())
    story.append(p("RÉFÉRENCES TECHNIQUES", st["h2"]))
    for project in PROJECTS:
        story.append(
            KeepTogether(
                [
                    rich(f"<b>{project['title']}</b><br/><font color='#667085'>{project['status']}</font>", st["h3"]),
                    *bullets(
                        [
                            f"Rôle: {project['role']}",
                            f"Technologies: {project['stack']}",
                            f"Périmètre: {project['scope']}",
                            f"Preuve: {project['evidence']}",
                            f"Lien: {project['url']}",
                        ],
                        st,
                    ),
                ]
            )
        )
    story.append(p("FORMATION ET CERTIFICATIONS", st["h2"]))
    story.extend(bullets([f"{a} | {b} | {c}" for a, b, c in EDUCATION], st))
    story.extend(bullets(CERTIFICATIONS, st))
    story.append(
        p(
            "Copies à joindre: les scans des deux diplômes et des trois certificats ne figuraient pas parmi les fichiers fournis.",
            st["callout"],
        )
    )
    story.append(p("RÔLE ATEA", st["h2"]))
    story.extend(
        bullets(
            [
                "Audit et industrialisation du code source LMS existant.",
                "Développement des rôles, cours, évaluations, progression, certificats et rapports.",
                "Tests, sécurité applicative, déploiement, documentation et transfert de compétences.",
                "Moodle reste un scénario comparatif; aucune expérience Moodle antérieure n'est revendiquée.",
            ],
            st,
        )
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def table(data, widths, header=True, font_size=7.7) -> Table:
    st = styles()
    converted = []
    for row_index, row in enumerate(data):
        row_style = ParagraphStyle(
            f"cell-{row_index}",
            parent=st["body"],
            fontName="AteaSans-Bold" if header and row_index == 0 else "AteaSans",
            fontSize=font_size,
            leading=font_size + 2.2,
            textColor=colors.white if header and row_index == 0 else colors.HexColor("#1D2939"),
            alignment=TA_LEFT,
        )
        converted.append([p(str(cell), row_style) for cell in row])
    tbl = Table(converted, colWidths=widths, repeatRows=1 if header else 0, hAlign="LEFT")
    commands = [
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("GRID", (0, 0), (-1, -1), 0.45, colors.HexColor("#C9D5DE")),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]
    if header:
        commands.append(("BACKGROUND", (0, 0), (-1, 0), NAVY))
        for r in range(1, len(data)):
            commands.append(("BACKGROUND", (0, r), (-1, r), LIGHT_GREY if r % 2 == 0 else colors.white))
    tbl.setStyle(TableStyle(commands))
    return tbl


def cover_story(st, title: str, subtitle: str, label: str) -> list:
    return [
        Spacer(1, 2.2 * cm),
        rich(f"<font color='#0E7C86'><b>{label}</b></font>", st["subtitle"]),
        Spacer(1, 0.5 * cm),
        p(title, st["title"]),
        Spacer(1, 0.35 * cm),
        p(subtitle, ParagraphStyle("cover-sub", parent=st["subtitle"], fontSize=14, leading=19)),
        Spacer(1, 1.1 * cm),
        Table(
            [
                [p("Candidat", st["small"]), rich("<b>Aymen Khaled</b>", st["body"])],
                [p("Rôle proposé", st["small"]), p(PROFILE["role"], st["body"])],
                [p("Soumissionnaire", st["small"]), rich("<b>Cabinet LE PLUS</b>", st["body"])],
                [p("Mission", st["small"]), p("Plateforme de formation en ligne et supports numériques ATEA", st["body"])],
                [p("Date", st["small"]), p("28 juin 2026", st["body"])],
            ],
            colWidths=[3.2 * cm, 12.8 * cm],
            style=TableStyle(
                [
                    ("BACKGROUND", (0, 0), (0, -1), LIGHT_BLUE),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#C9D5DE")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]
            ),
        ),
        Spacer(1, 1.2 * cm),
        p(
            "Document de preuve technique. Il ne remplace ni les références contractuelles du Cabinet LE PLUS, "
            "ni les copies officielles des diplômes et certificats.",
            st["callout"],
        ),
        PageBreak(),
    ]


def build_dossier_pdf(path: Path) -> None:
    st = styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=1.55 * cm,
        rightMargin=1.55 * cm,
        topMargin=1.45 * cm,
        bottomMargin=1.6 * cm,
        title="Dossier expert technique ATEA - Aymen Khaled",
        author="Aymen Khaled",
    )
    story = cover_story(
        st,
        "Dossier de compétences techniques et références",
        "Développement et industrialisation d'une plateforme LMS sur mesure pour l'ATEA",
        "PROJET HORIZON RECHERCHE - ATEA",
    )
    story.extend(
        [
            p("1. SYNTHÈSE DE POSITIONNEMENT", st["h1"]),
            p(
                "Aymen Khaled est proposé comme développeur Full-Stack responsable de la partie strictement technique. "
                "Il confirme disposer du code source d'un LMS similaire présenté dans la vidéo et les captures fournies. "
                "Le scénario principal consiste à auditer, adapter, compléter, tester, documenter et déployer cette base "
                "existante. Moodle est conservé uniquement comme scénario comparatif d'acquisition/configuration.",
                st["body"],
            ),
            p(
                "Positionnement honnête: Aymen n'est pas présenté comme expert Moodle ni comme expert en ingénierie "
                "pédagogique. Les contenus, études de cas et règles d'évaluation restent validés par les experts pédagogiques.",
                st["callout"],
            ),
            p("Responsabilités proposées", st["h2"]),
            *bullets(
                [
                    "Audit du code source, inventaire des modules existants et analyse des écarts avec le cahier des charges.",
                    "Industrialisation de l'architecture, des données, de la sécurité, des sauvegardes et du déploiement.",
                    "Développement des espaces par rôle, cours, contenus, évaluations, progression, certificats et rapports.",
                    "Intégration des contenus validés sans intervention dans leur conception pédagogique.",
                    "Tests multi-rôles, documentation, formation technique des administrateurs et support de démarrage.",
                ],
                st,
            ),
            p("2. CORRESPONDANCE AVEC LES BESOINS ATEA", st["h1"]),
        ]
    )
    matrix = [
        ["Besoin TDR", "Élément disponible", "Statut de preuve"],
        ["Cours et contenus centralisés", "Formulaires de cours/leçons, catalogue et page de cours", "Visible dans le démonstrateur"],
        ["Profils utilisateurs", "Espaces enseignant, apprenant et administrateur", "Visible; rôles ATEA à adapter"],
        ["Formation en ligne/hybride", "Cours, ressources, sessions en direct et outils collaboratifs", "Visible au niveau démonstrateur"],
        ["Évaluation des acquis", "Sélection cours/leçon et interface de génération de quiz", "Interface visible; résultats à tester"],
        ["Suivi des parcours", "Tableau de bord apprenant, progression et sessions", "Visible au niveau démonstrateur"],
        ["Administration et rapports", "Gestion utilisateurs et tableau analytique", "Visible; données de test non jointes"],
        ["Certificats", "Prévu dans l'architecture cible", "Non démontré dans les captures fournies"],
        ["Sécurité et performance", "Compétences techniques déclarées", "Rapports de tests à produire pendant la mission"],
        ["Autonomie ATEA", "Administration, guides et formation prévus", "Livrables futurs de la mission"],
    ]
    story.append(table(matrix, [5.0 * cm, 7.0 * cm, 4.1 * cm], font_size=7.2))
    story.append(PageBreak())
    story.append(p("3. FORMATION ET CERTIFICATIONS", st["h1"]))
    story.append(p("Diplômes déclarés dans le CV fourni", st["h2"]))
    story.extend(bullets([f"{a} - {b} - {c}" for a, b, c in EDUCATION], st))
    story.append(p("Certifications déclarées", st["h2"]))
    story.extend(bullets(CERTIFICATIONS, st))
    story.append(
        p(
            "État documentaire: aucune copie de diplôme, aucun fichier de certificat et aucune URL de vérification "
            "n'ont été fournis. Ces mentions peuvent figurer dans le CV, mais les justificatifs doivent être ajoutés "
            "avant l'envoi final au comité de sélection.",
            st["callout"],
        )
    )
    story.append(p("Compétences techniques mobilisables", st["h2"]))
    story.extend(
        bullets(
            [
                "React, Next.js, TypeScript, interfaces responsives et tableaux de bord.",
                "Node.js, Express, FastAPI, APIs REST/GraphQL et WebSockets.",
                "PostgreSQL, MongoDB, Redis et modélisation de données multi-utilisateurs.",
                "JWT/OAuth2, contrôle d'accès par rôles, validation et journalisation applicative.",
                "Docker, CI/CD, Linux, Nginx, stockage objet, supervision et sauvegardes.",
            ],
            st,
        )
    )
    story.append(p("4. RÉFÉRENCES TECHNIQUES", st["h1"]))
    for index, project in enumerate(PROJECTS, 1):
        project_data = [
            ["Statut", project["status"]],
            ["Période", project["period"]],
            ["Rôle", project["role"]],
            ["Technologies", project["stack"]],
            ["Périmètre", project["scope"]],
            ["Preuve disponible", project["evidence"]],
            ["Lien", project["url"]],
            ["Limite", project["limits"]],
        ]
        story.append(
            KeepTogether(
                [
                    p(f"4.{index} {project['title']}", st["h2"]),
                    table(project_data, [3.3 * cm, 12.8 * cm], header=False, font_size=7.5),
                    Spacer(1, 8),
                ]
            )
        )
    story.append(PageBreak())
    story.append(p("5. ARCHITECTURE CIBLE DU LMS SUR MESURE", st["h1"]))
    story.append(
        p(
            "Le développement ne consiste pas à coder les contenus pédagogiques. Les experts créent et valident "
            "les contenus; l'équipe technique fournit une interface d'administration permettant de structurer les "
            "modules, importer les ressources, configurer les évaluations et suivre les résultats.",
            st["body"],
        )
    )
    architecture = [
        ["Couche", "Responsabilité"],
        ["Interfaces React/TypeScript", "Espaces apprenant, formateur, administrateur, référent, expert et entité évaluée"],
        ["API Node.js", "Authentification, rôles, cours, progression, évaluations, certificats, rapports et intégrations"],
        ["PostgreSQL", "Données structurées, traçabilité, résultats, historique et relations métier"],
        ["Stockage objet", "PDF, vidéos, guides, études de cas, pièces jointes et attestations"],
        ["Services transverses", "Email, sauvegardes, audit, monitoring, sécurité, exports et IA optionnelle"],
        ["Exploitation", "Staging, production, CI/CD, tests de restauration, documentation et transfert"],
    ]
    story.append(table(architecture, [5.1 * cm, 11.0 * cm], font_size=7.5))
    story.append(p("Scénarios à présenter au comité", st["h2"]))
    scenarios = [
        ["Scénario", "Position dans l'offre", "Condition"],
        ["A - LMS sur mesure existant", "Recommandé", "Audit du code, propriété confirmée, industrialisation et tests d'acceptation"],
        ["B - Moodle", "Alternative d'acquisition/configuration", "Intervention d'un expert Moodle qualifié; Aymen ne revendique pas cette expertise"],
        ["C - SaaS tiers", "Option comparative", "Coût récurrent, dépendance fournisseur, localisation des données et réversibilité à évaluer"],
    ]
    story.append(table(scenarios, [4.2 * cm, 5.1 * cm, 6.8 * cm], font_size=7.2))
    story.append(p("6. DOSSIER DE PREUVE ET LIMITES", st["h1"]))
    story.extend(
        bullets(
            [
                "La vidéo courte de 1 min 30 s est le support principal recommandé; la vidéo de 5 min reste un lien complémentaire.",
                "L'annexe visuelle contient uniquement des captures réelles du démonstrateur, avec légendes de portée.",
                "Les références Strategy Navigator et Dashboard sont complémentaires, pas des missions LMS équivalentes.",
                "Les deux missions similaires obligatoires du prestataire doivent être fournies par Cabinet LE PLUS avec montant, date et destinataire.",
                "Les justificatifs de diplômes, certificats et expériences doivent être ajoutés dès qu'ils sont disponibles.",
            ],
            st,
        )
    )
    story.append(p("Portfolio et contacts", st["h2"]))
    story.append(
        table(
            [
                ["Portfolio", PROFILE["portfolio"]],
                ["LinkedIn", PROFILE["linkedin"]],
                ["GitHub", PROFILE["github"]],
                ["Email", PROFILE["email"]],
            ],
            [3.2 * cm, 12.9 * cm],
            header=False,
        )
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


SCREENSHOTS = [
    (
        "02_create_new_course_modal.png",
        "Création d'un cours et d'une leçon",
        "Éléments visibles: titre, durée, description, URL ou fichier vidéo et structure de leçon.",
        "Portée: preuve d'interface du démonstrateur; la persistance et la validation serveur doivent être testées.",
    ),
    (
        "04_course_recording_preview_area.png",
        "Studio de production pédagogique",
        "Éléments visibles: enregistrement, audio, tableau blanc, éditeur de code, carte mentale, plan de leçon et médias.",
        "Portée: inventaire fonctionnel de l'interface; chaque outil doit être vérifié séparément en recette.",
    ),
    (
        "05_tutoring_sessions_list.png",
        "Gestion des sessions de tutorat",
        "Éléments visibles: liste, date, durée, capacité, modification, suppression et création de session.",
        "Portée: flux directement pertinent pour la formation synchrone/hybride.",
    ),
    (
        "09_live_session_webcam.png",
        "Session en direct connectée",
        "Éléments visibles: état connecté, chronomètre, rôle hôte et flux webcam.",
        "Portée: preuve visuelle d'une session locale; les tests multi-participants et de charge restent à produire.",
    ),
    (
        "12_course_lesson_selection.png",
        "Sélection d'un cours et d'une leçon",
        "Éléments visibles: navigation par cours puis leçon avant génération d'une activité ou d'un quiz.",
        "Portée: flux utile pour la création d'évaluations contextualisées.",
    ),
    (
        "16_courses_catalog.png",
        "Catalogue de cours",
        "Éléments visibles: recherche, filtres, catégories, durée, nombre d'apprenants et inscription.",
        "Portée: les fonctions de paiement visibles dans le prototype ne sont pas nécessaires pour le périmètre ATEA.",
    ),
    (
        "17_course_detail_page.png",
        "Détail d'un cours et parcours",
        "Éléments visibles: description, durée, programme, aperçu vidéo, formateur et avis.",
        "Portée: la version ATEA remplacera l'achat par l'affectation institutionnelle des parcours.",
    ),
]


def screenshot_bytes(path: Path) -> io.BytesIO:
    with PILImage.open(path) as image:
        image = image.convert("RGB")
        width, height = image.size
        cropped = image.crop((0, 66, width, height - 38))
        stream = io.BytesIO()
        cropped.save(stream, format="JPEG", quality=92, optimize=True)
        stream.seek(0)
        return stream


def build_visual_annex(path: Path) -> None:
    register_fonts()
    page_w, page_h = landscape(A4)
    c = canvas.Canvas(str(path), pagesize=(page_w, page_h), pageCompression=1)
    c.setTitle("Annexe visuelle LMS ATEA - Aymen Khaled")
    c.setAuthor("Aymen Khaled")
    c.setFillColor(NAVY)
    c.rect(0, 0, page_w, page_h, fill=1, stroke=0)
    c.setFillColor(colors.white)
    c.setFont("AteaSans-Bold", 26)
    c.drawString(2.0 * cm, page_h - 3.0 * cm, "Annexe visuelle du démonstrateur LMS")
    c.setFont("AteaSans", 14)
    c.setFillColor(colors.HexColor("#BFE5EA"))
    c.drawString(2.0 * cm, page_h - 4.0 * cm, "Captures réelles - environnement local - mai 2025")
    c.setFillColor(colors.white)
    c.setFont("AteaSans", 10.5)
    text = c.beginText(2.0 * cm, page_h - 6.0 * cm)
    text.setLeading(16)
    for line in [
        "Cette annexe illustre une base logicielle existante similaire au besoin ATEA.",
        "Elle ne constitue pas une attestation de réception client ni un rapport de tests.",
        "Les légendes distinguent ce qui est visible de ce qui reste à vérifier en recette.",
    ]:
        text.textLine(line)
    c.drawText(text)
    c.setFillColor(colors.HexColor("#BFE5EA"))
    c.setFont("AteaSans", 9)
    c.drawString(2.0 * cm, 1.5 * cm, "Cabinet LE PLUS - Dossier ATEA - Aymen Khaled")
    c.showPage()

    for index, (filename, title, evidence, scope) in enumerate(SCREENSHOTS, 1):
        source = SCREENSHOT_DIR / filename
        if not source.exists():
            raise FileNotFoundError(f"Capture LMS manquante: {source}")
        c.setFillColor(colors.white)
        c.rect(0, 0, page_w, page_h, fill=1, stroke=0)
        c.setFillColor(NAVY)
        c.rect(0, page_h - 2.3 * cm, page_w, 2.3 * cm, fill=1, stroke=0)
        c.setFillColor(colors.white)
        c.setFont("AteaSans-Bold", 16)
        c.drawString(1.4 * cm, page_h - 1.45 * cm, f"{index}. {title}")
        stream = screenshot_bytes(source)
        with PILImage.open(stream) as img:
            iw, ih = img.size
        max_w = page_w - 2.8 * cm
        max_h = page_h - 7.1 * cm
        scale = min(max_w / iw, max_h / ih)
        draw_w, draw_h = iw * scale, ih * scale
        x = (page_w - draw_w) / 2
        y = 3.9 * cm + (max_h - draw_h) / 2
        stream.seek(0)
        c.drawImage(ImageReader(stream), x, y, width=draw_w, height=draw_h, preserveAspectRatio=True, mask="auto")
        c.setFillColor(LIGHT_BLUE)
        c.roundRect(1.4 * cm, 1.0 * cm, page_w - 2.8 * cm, 2.25 * cm, 0.12 * cm, fill=1, stroke=0)
        caption_body = ParagraphStyle(
            "AnnexCaptionBody",
            fontName="AteaSans",
            fontSize=8.1,
            leading=9.4,
            textColor=colors.HexColor("#1D2939"),
        )
        caption_label = ParagraphStyle(
            "AnnexCaptionLabel",
            fontName="AteaSans-Bold",
            fontSize=8.5,
            leading=9.4,
            textColor=NAVY,
        )
        caption_scope_label = ParagraphStyle(
            "AnnexCaptionScopeLabel",
            fontName="AteaSans-Bold",
            fontSize=8.5,
            leading=9.4,
            textColor=AMBER,
        )
        caption_table = Table(
            [
                [Paragraph("ÉLÉMENTS VISIBLES", caption_label), Paragraph(evidence, caption_body)],
                [Paragraph("PORTÉE", caption_scope_label), Paragraph(scope, caption_body)],
            ],
            colWidths=[3.25 * cm, page_w - 6.35 * cm],
            rowHeights=[0.82 * cm, 0.82 * cm],
        )
        caption_table.setStyle(
            TableStyle(
                [
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ]
            )
        )
        caption_table.wrapOn(c, page_w - 3.1 * cm, 1.64 * cm)
        caption_table.drawOn(c, 1.7 * cm, 1.30 * cm)
        c.setFillColor(MID_GREY)
        c.setFont("AteaSans", 7.5)
        c.drawRightString(page_w - 1.4 * cm, 0.55 * cm, f"Capture {index}/{len(SCREENSHOTS)} - {filename}")
        c.showPage()
    c.save()


def draw_box(c, x, y, w, h, title, lines, fill_color) -> None:
    c.setFillColor(fill_color)
    c.setStrokeColor(colors.HexColor("#B8C7D1"))
    c.roundRect(x, y, w, h, 0.12 * cm, fill=1, stroke=1)
    c.setFillColor(NAVY)
    c.setFont("AteaSans-Bold", 10)
    c.drawString(x + 0.3 * cm, y + h - 0.55 * cm, title)
    c.setFont("AteaSans", 7.5)
    c.setFillColor(colors.HexColor("#344054"))
    text = c.beginText(x + 0.3 * cm, y + h - 1.0 * cm)
    text.setLeading(10)
    for line in lines:
        text.textLine(line)
    c.drawText(text)


def arrow(c, x1, y1, x2, y2) -> None:
    c.setStrokeColor(BLUE)
    c.setFillColor(BLUE)
    c.setLineWidth(1.5)
    c.line(x1, y1, x2, y2)
    c.line(x2, y2, x2 - 6, y2 + 3)
    c.line(x2, y2, x2 - 6, y2 - 3)


def build_architecture_pdf(path: Path) -> None:
    register_fonts()
    page_w, page_h = landscape(A4)
    c = canvas.Canvas(str(path), pagesize=(page_w, page_h), pageCompression=1)
    c.setTitle("Architecture LMS sur mesure ATEA")
    c.setAuthor("Aymen Khaled")
    c.setFillColor(NAVY)
    c.setFont("AteaSans-Bold", 22)
    c.drawString(1.5 * cm, page_h - 1.7 * cm, "Architecture proposée - LMS sur mesure ATEA")
    c.setFillColor(MID_GREY)
    c.setFont("AteaSans", 9)
    c.drawString(1.5 * cm, page_h - 2.25 * cm, "Scénario principal: audit, adaptation et industrialisation du code source LMS existant")

    y_top = page_h - 5.0 * cm
    box_h = 2.3 * cm
    draw_box(c, 1.5 * cm, y_top, 4.4 * cm, box_h, "UTILISATEURS", ["Apprenant - Formateur", "Admin - Référent - Expert"], LIGHT_BLUE)
    draw_box(c, 7.0 * cm, y_top, 5.0 * cm, box_h, "INTERFACES REACT", ["Portail par rôle", "Cours - quiz - tableaux de bord"], colors.HexColor("#E7F6F1"))
    draw_box(c, 13.1 * cm, y_top, 5.0 * cm, box_h, "API NODE.JS", ["RBAC - progression - rapports", "Intégrations - notifications"], colors.HexColor("#FFF4E5"))
    draw_box(c, 19.2 * cm, y_top, 4.7 * cm, box_h, "DONNÉES", ["PostgreSQL", "Stockage fichiers protégé"], colors.HexColor("#F4EBFF"))
    draw_box(c, 24.9 * cm, y_top, 3.3 * cm, box_h, "SERVICES", ["Email - IA", "Monitoring"], colors.HexColor("#FDECEC"))
    arrow(c, 5.9 * cm, y_top + box_h / 2, 7.0 * cm, y_top + box_h / 2)
    arrow(c, 12.0 * cm, y_top + box_h / 2, 13.1 * cm, y_top + box_h / 2)
    arrow(c, 18.1 * cm, y_top + box_h / 2, 19.2 * cm, y_top + box_h / 2)
    arrow(c, 23.9 * cm, y_top + box_h / 2, 24.9 * cm, y_top + box_h / 2)

    draw_box(c, 2.0 * cm, 5.3 * cm, 8.0 * cm, 3.0 * cm, "CONTENU SANS CODAGE", ["Les experts pédagogiques utilisent", "l'administration pour créer modules,", "ressources, cas, exercices et quiz."], colors.HexColor("#E7F6F1"))
    draw_box(c, 11.0 * cm, 5.3 * cm, 8.0 * cm, 3.0 * cm, "DÉVELOPPEMENT AYMEN", ["Code des fonctions, APIs, rôles,", "rapports, tests, déploiement,", "documentation et maintenance."], LIGHT_BLUE)
    draw_box(c, 20.0 * cm, 5.3 * cm, 7.6 * cm, 3.0 * cm, "EXPLOITATION ATEA", ["Staging et production", "Sauvegardes et restauration", "Formation et transfert"], colors.HexColor("#FFF4E5"))
    c.setFillColor(NAVY)
    c.setFont("AteaSans-Bold", 11)
    c.drawString(1.5 * cm, 3.9 * cm, "Alternative de comparaison")
    c.setFont("AteaSans", 9)
    c.setFillColor(colors.HexColor("#344054"))
    c.drawString(1.5 * cm, 3.35 * cm, "Moodle peut être présenté comme scénario d'acquisition/configuration, sous la responsabilité d'un expert Moodle qualifié.")
    c.setFillColor(RED)
    c.setFont("AteaSans-Bold", 9)
    c.drawString(1.5 * cm, 2.75 * cm, "Aucune expérience Moodle antérieure n'est revendiquée pour Aymen Khaled.")
    c.setFillColor(MID_GREY)
    c.setFont("AteaSans", 7.5)
    c.drawRightString(page_w - 1.5 * cm, 0.75 * cm, "Architecture cible à valider pendant le cadrage technique")
    c.save()


def mehdi_content() -> dict:
    return {
        "role_line": (
            "Aymen Khaled, Développeur Full-Stack - Responsable du développement technique de la plateforme LMS;"
        ),
        "scope": (
            "La partie développement couvre l'audit du code source LMS existant, l'analyse des écarts avec les besoins ATEA, "
            "l'adaptation fonctionnelle et graphique, l'industrialisation de l'architecture, le développement des modules "
            "complémentaires, l'intégration des contenus validés, les tests, le déploiement, la documentation et le transfert "
            "de compétences. La scénarisation pédagogique, la rédaction des contenus et l'expertise en évaluation restent "
            "sous la responsabilité des experts pédagogiques et du chef de mission."
        ),
        "approach": (
            "Le scénario technique recommandé consiste à adapter et industrialiser une base logicielle LMS existante dont "
            "Aymen Khaled dispose du code source. Cette approche réduit le délai de réalisation tout en permettant une "
            "personnalisation complète selon les rôles, parcours, référentiels, règles d'évaluation et exigences de traçabilité "
            "de l'ATEA. La base fera l'objet d'un audit initial portant sur la qualité du code, la sécurité, le modèle de données, "
            "les dépendances, les licences, les performances et la couverture fonctionnelle. La mise en production sera soumise "
            "à des tests d'acceptation formalisés. Conformément au TDR, un scénario alternatif d'acquisition et de configuration "
            "de Moodle sera présenté dans le rapport comparatif. Moodle ne constitue pas l'engagement principal d'Aymen et toute "
            "intervention Moodle devra être confiée à un expert disposant d'une expérience démontrable sur cette technologie."
        ),
        "content_workflow": (
            "La création de contenu ne nécessite pas de programmation au quotidien. Les experts pédagogiques et administrateurs "
            "ATEA utiliseront l'interface d'administration pour créer les cours, structurer les modules, importer des PDF ou des "
            "vidéos, ajouter des études de cas, configurer les quiz, définir les seuils de réussite et publier les parcours. Le "
            "développeur intervient pour construire, sécuriser et maintenir ces fonctionnalités, et non pour rédiger le contenu pédagogique."
        ),
    }


STACK_ROWS = [
    ["Composant", "Choix proposé", "Justification"],
    ["Base existante", "Code source LMS détenu par Aymen", "Réduction du délai après audit technique et juridique"],
    ["Frontend", "React + TypeScript", "Interfaces responsives par rôle et administration des contenus"],
    ["Backend", "Node.js + API REST", "Règles métier, rôles, progression, rapports et intégrations"],
    ["Base de données", "PostgreSQL cible", "Intégrité des parcours, évaluations, certificats et traçabilité"],
    ["Fichiers", "Stockage objet compatible S3", "PDF, vidéos, guides, exercices et attestations"],
    ["Sécurité", "TLS, RBAC, validation, audit, sauvegardes", "Protection des accès et reprise après incident"],
    ["Déploiement", "Staging + production, Docker, Nginx, CI/CD", "Recette contrôlée et mises à jour reproductibles"],
    ["Alternative", "Moodle configuré par un expert Moodle", "Scénario comparatif d'acquisition demandé par le TDR"],
]

SCHEDULE_ROWS = [
    ["Phase", "Durée", "Validation"],
    ["Cadrage, audit du code et scénarios", "2 semaines", "Rapport d'audit et choix du scénario"],
    ["Architecture et cahier des charges", "2 semaines", "Validation du comité de pilotage"],
    ["Adaptation et développement", "6 à 8 semaines", "Démonstrations par lot"],
    ["Intégration des contenus", "3 à 4 semaines", "Recette pédagogique"],
    ["Tests, sécurité et corrections", "2 semaines", "Rapport de tests et PV de recette"],
    ["Production, guides et transfert", "2 semaines", "Mise en service et formation administrateurs"],
]


def build_mehdi_docx(path: Path) -> None:
    content = mehdi_content()
    doc = make_doc(path, "Corrections proposition ATEA - Cabinet LE PLUS", "Aymen Khaled")
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("CORRECTIONS PRÊTES À INTÉGRER DANS LA PROPOSITION ATEA")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(11, 46, 79)
    doc.add_paragraph("Cabinet LE PLUS | Version de travail du 28/06/2026").alignment = WD_ALIGN_PARAGRAPH.CENTER

    docx_heading(doc, "1. Remplacement dans la lettre de soumission", 1)
    doc.add_paragraph("Remplacer la ligne actuelle concernant Aymen par:")
    q = doc.add_paragraph()
    qr = q.add_run(content["role_line"])
    qr.bold = True
    qr.font.color.rgb = RGBColor(14, 124, 134)
    doc.add_paragraph(
        "Aymen est un expert technique additionnel. Il ne doit pas être comptabilisé comme l'un des deux experts "
        "obligatoires en pédagogie numérique."
    )

    docx_heading(doc, "2. Périmètre de la partie développement", 1)
    doc.add_paragraph(content["scope"])
    docx_heading(doc, "3. Approche technique recommandée", 1)
    doc.add_paragraph(content["approach"])
    docx_heading(doc, "4. Relation entre contenu pédagogique et développement", 1)
    doc.add_paragraph(content["content_workflow"])

    docx_heading(doc, "5. Stack technique proposée", 1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for i, value in enumerate(STACK_ROWS[0]):
        tbl.rows[0].cells[i].text = value
    for row in STACK_ROWS[1:]:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    docx_heading(doc, "6. Chronogramme technique", 1)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for i, value in enumerate(SCHEDULE_ROWS[0]):
        tbl.rows[0].cells[i].text = value
    for row in SCHEDULE_ROWS[1:]:
        cells = tbl.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    docx_heading(doc, "7. Offre en tiroirs recommandée", 1)
    for item in [
        "Tiroir A - Cadrage, audit du code, comparaison des scénarios et cahier des charges.",
        "Tiroir B - Adaptation et développement de la plateforme LMS sur mesure.",
        "Tiroir C - Intégration des contenus, tests, mise en production, guides et formation.",
        "Option 1 - Hébergement, supervision, sauvegardes et maintenance pendant douze mois.",
        "Option 2 - Fonctionnalités IA, classes virtuelles avancées ou micro-certifications.",
        "Option 3 - Acquisition/configuration Moodle par un expert qualifié, si ce scénario est retenu.",
    ]:
        docx_bullet(doc, item)

    docx_heading(doc, "8. Éléments à supprimer ou corriger", 1)
    for item in [
        "Supprimer tous les champs génériques relatifs au nom du Cabinet, à la ville, à la date et au nom du signataire, ainsi que la mention Document technique interne.",
        "Corriger plateforme web modernes en plateforme web moderne.",
        "Ne pas annoncer une licence Moodle si l'offre porte sur Moodle LMS open source; chiffrer plutôt hébergement, support et plugins commerciaux éventuels.",
        "Ne pas déclarer Aymen expert Moodle, expert pédagogique ou titulaire de cinq années d'expérience pédagogique.",
        "Ne pas appeler les projets portfolio missions similaires du prestataire sans contrat, montant, date et destinataire.",
        "Ajouter les critères d'acceptation, sauvegardes, restauration, journalisation, sécurité, performance, réversibilité et maintenance.",
    ]:
        docx_bullet(doc, item)
    doc.save(path)


def build_mehdi_pdf(path: Path) -> None:
    st = styles()
    content = mehdi_content()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.45 * cm,
        bottomMargin=1.6 * cm,
        title="Corrections proposition ATEA - Cabinet LE PLUS",
        author="Aymen Khaled",
    )
    story = [
        p("Corrections prêtes à intégrer dans la proposition ATEA", st["title"]),
        p("Cabinet LE PLUS | Version de travail du 28/06/2026", st["subtitle"]),
        p("1. Lettre de soumission", st["h1"]),
        p("Remplacer la ligne actuelle concernant Aymen par:", st["body"]),
        p(content["role_line"], st["callout"]),
        p(
            "Aymen est un expert technique additionnel. Il ne doit pas être comptabilisé comme l'un des deux experts obligatoires en pédagogie numérique.",
            st["body"],
        ),
        p("2. Périmètre développement", st["h1"]),
        p(content["scope"], st["body"]),
        p("3. Approche technique recommandée", st["h1"]),
        p(content["approach"], st["body"]),
        p("4. Contenu pédagogique et développement", st["h1"]),
        p(content["content_workflow"], st["body"]),
        PageBreak(),
        p("5. Stack technique", st["h1"]),
        table(STACK_ROWS, [3.1 * cm, 5.1 * cm, 8.2 * cm], font_size=7.1),
        p("6. Chronogramme", st["h1"]),
        table(SCHEDULE_ROWS, [6.5 * cm, 2.6 * cm, 7.3 * cm], font_size=7.2),
        p("7. Offre en tiroirs", st["h1"]),
        *bullets(
            [
                "Tiroir A - Cadrage, audit, scénarios et cahier des charges.",
                "Tiroir B - Adaptation et développement du LMS sur mesure.",
                "Tiroir C - Contenus, tests, production, guides et formation.",
                "Options - Hébergement/maintenance, IA/micro-certifications, ou Moodle par un expert qualifié.",
            ],
            st,
        ),
        p("8. Corrections indispensables", st["h1"]),
        *bullets(
            [
                "Supprimer les placeholders et la mention document technique interne.",
                "Corriger plateforme web modernes en plateforme web moderne.",
                "Ne pas présenter Aymen comme expert Moodle ou expert en pédagogie numérique.",
                "Ne pas présenter les projets portfolio comme références contractuelles du Cabinet LE PLUS.",
                "Ajouter sécurité, restauration, acceptation, réversibilité et maintenance.",
            ],
            st,
        ),
    ]
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def build_checklist(path: Path) -> None:
    content = """# Pièces justificatives manquantes - Dossier ATEA

Date de contrôle: 28/06/2026

## Priorité critique avant l'envoi

- [ ] Copie de la Licence en Big Data et Analyse de Données - ISIMS, 2025.
- [ ] Copie du Diplôme en Génie Logiciel - ESPS, 2022.
- [ ] Certificat JavaScript Algorithms - freeCodeCamp, avec URL de vérification si disponible.
- [ ] Certificat Responsive Web Design - freeCodeCamp, avec URL de vérification si disponible.
- [ ] Certificat Young Entrepreneur Rookies.
- [ ] Attestation de stage ou lettre de référence Aziin Engineering Solution concernant le LMS.
- [ ] Attestation employeur Everything to Gain précisant le rôle sur Strategy Navigator/JourneyAI/SaleSide.
- [ ] Déclaration signée de propriété et de droit de réutilisation du code source LMS pour la mission ATEA.

## Pièces obligatoires relevant du Cabinet LE PLUS

- [ ] Au moins deux missions similaires du prestataire, avec montant, date et destinataire public ou privé.
- [ ] Preuves contractuelles ou attestations de bonne exécution pour ces missions.
- [ ] Justificatif d'exercice du Cabinet depuis au moins trois ans.
- [ ] CV et diplômes du chef de mission.
- [ ] CV, diplômes BAC+5 et preuves de cinq années d'expérience des deux experts en pédagogie numérique.
- [ ] Preuves de trois missions similaires pour chaque profil concerné lorsque le TDR l'exige.
- [ ] Moyens humains et matériels du Cabinet.
- [ ] Offre financière en tiroirs avec détail par poste, fournisseurs, imprévus et provision pour risques.
- [ ] Note conceptuelle finale de trois à cinq pages.

## Contrôle de cohérence

- [ ] Aymen est présenté comme Développeur Full-Stack, pas comme expert Moodle ou expert pédagogique.
- [ ] EduNova est présenté comme démonstrateur de stage tant qu'aucune attestation plus forte n'est jointe.
- [ ] Strategy Navigator et Dashboard sont présentés comme références techniques complémentaires.
- [ ] Les captures contenant emails personnels, données de test ou défauts visibles ne sont pas envoyées séparément.
- [ ] La vidéo courte de 1 min 30 s est utilisée en priorité; la vidéo de 5 min reste un lien complémentaire.
"""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")


def build_readme(path: Path) -> None:
    content = """# Pack d'envoi ATEA - Aymen Khaled

## Fichiers à transmettre à Mehdi

1. `pdf/Dossier_Expert_Technique_Aymen_Khaled_ATEA.pdf`
2. `pdf/CV_Aymen_Khaled_ATEA_FR.pdf`
3. `pdf/Annexe_Visuelle_Demonstrateur_LMS_ATEA.pdf`
4. `pdf/Architecture_LMS_Sur_Mesure_ATEA.pdf`
5. `docx/Corrections_Proposition_Mehdi_ATEA.docx`
6. Les scans réels des diplômes, certificats et attestations listés dans le checklist.

## Démonstration recommandée

- Envoyer la vidéo professionnelle de 1 min 30 s avec le dossier.
- Mettre la vidéo de 5 min sur un lien privé/non répertorié comme preuve complémentaire.
- Ne pas envoyer le ZIP brut des 25 captures.

## Positionnement à conserver

**Aymen Khaled - Développeur Full-Stack - Responsable du développement technique de la plateforme LMS**

Le scénario principal est l'adaptation et l'industrialisation du code source LMS existant appartenant à Aymen. Moodle est une alternative comparative et non une compétence antérieure revendiquée.

## Avertissement documentaire

Les diplômes et certifications sont listés dans le CV fourni, mais leurs copies n'étaient pas présentes parmi les fichiers analysés. Aucun faux justificatif n'a été généré.
"""
    path.write_text(content, encoding="utf-8")


def assemble_zip() -> None:
    if FINAL_ZIP.exists():
        FINAL_ZIP.unlink()
    with zipfile.ZipFile(FINAL_ZIP, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for relative in required_artifacts():
            source = OUTPUT_ROOT / relative
            if not source.exists():
                raise FileNotFoundError(f"Livrable manquant: {source}")
            archive.write(source, arcname=relative.replace("\\", "/"))


def prepare_output() -> None:
    for directory in [PDF_DIR, DOCX_DIR, CHECKLIST_DIR]:
        directory.mkdir(parents=True, exist_ok=True)
    for relative in required_artifacts():
        target = OUTPUT_ROOT / relative
        if target.exists():
            target.unlink()


def main() -> None:
    prepare_output()
    build_cv_pdf(PDF_DIR / "CV_Aymen_Khaled_ATEA_FR.pdf")
    build_cv_docx(DOCX_DIR / "CV_Aymen_Khaled_ATEA_FR.docx")
    build_dossier_pdf(PDF_DIR / "Dossier_Expert_Technique_Aymen_Khaled_ATEA.pdf")
    build_visual_annex(PDF_DIR / "Annexe_Visuelle_Demonstrateur_LMS_ATEA.pdf")
    build_architecture_pdf(PDF_DIR / "Architecture_LMS_Sur_Mesure_ATEA.pdf")
    build_mehdi_docx(DOCX_DIR / "Corrections_Proposition_Mehdi_ATEA.docx")
    build_mehdi_pdf(PDF_DIR / "Corrections_Proposition_Mehdi_ATEA.pdf")
    build_checklist(CHECKLIST_DIR / "Pieces_Justificatives_Manquantes_ATEA.md")
    build_readme(OUTPUT_ROOT / "README_ENVOI.md")
    assemble_zip()
    print(f"Pack généré: {FINAL_ZIP}")


if __name__ == "__main__":
    main()
