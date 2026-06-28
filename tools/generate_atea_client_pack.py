from __future__ import annotations

import hashlib
import html
import io
import shutil
import zipfile
from pathlib import Path

from PIL import Image as PILImage
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image,
    KeepTogether,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
WORKSPACE = ROOT.parent
OUTPUT_ROOT = WORKSPACE / "output"
CLIENT_DIR = OUTPUT_ROOT / "client_final"
EDITABLE_DIR = OUTPUT_ROOT / "editable"
CLIENT_CV = CLIENT_DIR / "01_CV_Aymen_Khaled_Developpeur_Full_Stack_LMS_FR.pdf"
CLIENT_PORTFOLIO = CLIENT_DIR / "02_Portfolio_Technique_LMS_Aymen_Khaled.pdf"
CLIENT_DIPLOMA = CLIENT_DIR / "03_Justificatifs" / "Diplome_Licence_Big_Data_Analyse_Donnees.png"
SOURCE_DIPLOMA = Path(r"C:\Users\khale\Downloads\diplome.PNG")
CLIENT_ZIP = OUTPUT_ROOT / "ATEA_Aymen_Khaled_Dossier_Client_Final.zip"
EDITABLE_CV = EDITABLE_DIR / "CV_Aymen_Khaled_Developpeur_Full_Stack_LMS_FR.docx"
SCREENSHOT_DIR = (
    WORKSPACE
    / "tmp"
    / "evidence_analysis"
    / "lms_video_screenshots_and_analysis"
    / "screenshots_by_section"
)

NAVY = colors.HexColor("#0B2E4F")
BLUE = colors.HexColor("#1769AA")
TEAL = colors.HexColor("#0E7C86")
LIGHT_BLUE = colors.HexColor("#EAF3F8")
LIGHT_GREY = colors.HexColor("#F4F6F8")
MID_GREY = colors.HexColor("#667085")
TEXT = colors.HexColor("#1D2939")
WHITE = colors.white


PROFILE = {
    "name": "Aymen Khaled",
    "title": "Développeur Full-Stack | Plateformes SaaS, LMS et intégrations IA",
    "location": "Mahdia, Tunisie",
    "phone": "+216 26 286 045",
    "email": "khaledaymen850@gmail.com",
    "portfolio": "https://aymen-kh.vercel.app",
    "linkedin": "https://linkedin.com/in/aymen-khaled-652724236",
    "github": "https://github.com/aymenkhaled",
}

SUMMARY = (
    "Développeur Full-Stack spécialisé dans la conception de plateformes SaaS, d'applications web et mobile "
    "et d'intégrations d'intelligence artificielle. Expérience concrète des architectures multi-tenant, du contrôle "
    "d'accès par rôles, des API, des fonctionnalités temps réel, de la facturation en ligne et des chaînes CI/CD. "
    "Compétences directement mobilisables pour l'audit, l'adaptation et l'industrialisation d'une plateforme LMS."
)

SKILLS = [
    ("Langages", "JavaScript (ES6+), TypeScript, Python, C#, PHP, Java, C++"),
    ("Front-end", "React, Next.js, React Native, Redux, Tailwind CSS, HTML5, CSS3, interfaces adaptatives"),
    ("Back-end", "Node.js, Express, FastAPI, Django, Laravel, API REST, GraphQL, WebSockets, microservices"),
    ("Données", "PostgreSQL, MongoDB, SQL Server, Firebase, Redis, Supabase"),
    ("IA et LLM", "OpenAI GPT-4o, Groq Llama, Whisper, LangChain, ingénierie de prompts, agents IA"),
    ("DevOps et qualité", "Docker, Jenkins, GitHub Actions, CI/CD, Linux, Nginx, AWS S3/EC2, Jest, pytest, Postman"),
    ("Intégrations", "HubSpot, Slack, Asana, Basecamp, Stripe, PayPal, WhatsApp, OAuth2, webhooks"),
    ("Méthodes", "Agile, Scrum, Kanban, Jira, revue de code, documentation technique"),
]

EXPERIENCE = [
    {
        "role": "Développeur Web Full-Stack",
        "company": "Everything to Gain",
        "location": "À distance",
        "period": "Février 2025 - présent",
        "bullets": [
            "Conception et développement de Strategy Navigator, plateforme SaaS multi-tenant avec isolation des données et contrôle d'accès par rôles.",
            "Intégration de plus de 65 services d'IA derrière une couche API unifiée, avec traitement par lots et cache Redis.",
            "Mise en place de la facturation Stripe et PayPal, des webhooks d'activation et de trois niveaux d'abonnement.",
            "Intégration bidirectionnelle de Slack, Asana et Basecamp via OAuth et webhooks; industrialisation des déploiements avec Docker et Jenkins.",
            "Développement de JourneyAI: plus de dix assistants IA, chat WebSocket, transcription Whisper, routage OpenAI/Groq et facturation à l'usage.",
            "Développement de SaleSide AI, RxCare / Deep Analyzer Suite et Integrity Solve pour les domaines vente, santé et conformité AML/CTF.",
        ],
    },
    {
        "role": "Développeur Full-Stack - stage",
        "company": "Aziin Engineering Solution",
        "location": "Tunisie",
        "period": "2024",
        "bullets": [
            "Développement d'une plateforme e-learning en stack MERN avec assistant IA et classe virtuelle temps réel via WebSockets.",
            "Gestion dynamique des contenus, authentification JWT et optimisation des requêtes par indexation et chargement différé.",
        ],
    },
    {
        "role": "Développeur mobile - stage",
        "company": "SAC Marquage",
        "location": "Tunisie",
        "period": "Septembre 2024 - octobre 2024",
        "bullets": [
            "Application React Native / Expo connectée à des équipements RFID par Bluetooth et NFC, avec lecture, écriture et synchronisation des tags.",
            "Développement d'API Django REST pour la gestion des équipements et la synchronisation des états.",
        ],
    },
    {
        "role": "Développeur Web Full-Stack - stage",
        "company": "Proged",
        "location": "Tunisie",
        "period": "Juillet 2022 - août 2022",
        "bullets": [
            "Développement d'une application e-commerce avec back-end .NET / C#, front-end React, SQL Server et MongoDB.",
            "Réalisation des parcours catalogue, commande, transaction et paiement.",
        ],
    },
]

PROJECTS = [
    (
        "Job Scraper & Outreach Automation",
        "Python, MERN, microservices, OpenAI",
        "Collecte d'offres sur plus de 40 plateformes, enrichissement des contacts et génération d'emails personnalisés par IA.",
    ),
    (
        "AutoSEO",
        "React, FastAPI, PostgreSQL, Redis, Supabase",
        "Plateforme d'audit SEO, regroupement des causes racines, priorisation des corrections et création de propositions de modification GitHub.",
    ),
    (
        "AutoTN",
        "Expo, React Native, Express, PostgreSQL",
        "Marketplace trilingue pour véhicules et pièces avec diagnostic IA, décodage VIN, estimation des prix et parcours WhatsApp.",
    ),
    (
        "Outreach Personalization Engine",
        "Python, OpenAI, Node.js",
        "Chaîne d'enrichissement des prospects et de génération de messages contextualisés à partir des signaux entreprise et contact.",
    ),
    (
        "E-Commerce Dashboard",
        "React, Node.js, PostgreSQL",
        "Tableau de bord centralisé pour l'analyse d'activité, la gestion des stocks et le suivi des commandes.",
    ),
]

EDUCATION = [
    ("Licence en Big Data et Analyse de Données", "ISIMS, Sfax, Tunisie", "2022 - 2025"),
    ("Diplôme en Génie Logiciel", "ESPS, Tunisie", "2022"),
]

CERTIFICATIONS = [
    "JavaScript Algorithms - freeCodeCamp",
    "Responsive Web Design - freeCodeCamp",
    "Young Entrepreneur Rookies",
]

LANGUAGES = [
    "Arabe: langue maternelle",
    "Anglais: niveau avancé, pratique professionnelle",
    "Français: niveau intermédiaire B1",
]

SCREENSHOTS = [
    (
        "02_create_new_course_modal.png",
        "Création structurée des cours et des leçons",
        "Saisie des informations pédagogiques, durée, description, ressources vidéo et organisation des leçons.",
    ),
    (
        "04_course_recording_preview_area.png",
        "Studio de production pédagogique",
        "Enregistrement, audio, tableau blanc, éditeur de code, carte mentale, plan de leçon et outils multimédias.",
    ),
    (
        "05_tutoring_sessions_list.png",
        "Planification des sessions de tutorat",
        "Création, modification et suivi des sessions synchrones avec date, durée et capacité d'accueil.",
    ),
    (
        "09_live_session_webcam.png",
        "Classe virtuelle en temps réel",
        "Espace de session connecté avec rôle hôte, chronomètre, vidéo et commandes de session.",
    ),
    (
        "12_course_lesson_selection.png",
        "Évaluations contextualisées par cours",
        "Sélection du cours et de la leçon avant la création d'une activité ou d'un quiz assisté par IA.",
    ),
    (
        "16_courses_catalog.png",
        "Catalogue de formation",
        "Recherche, filtres, catégories, durée, nombre d'apprenants et accès aux parcours disponibles.",
    ),
    (
        "17_course_detail_page.png",
        "Présentation détaillée d'un parcours",
        "Description, programme, durée, aperçu vidéo, informations formateur et contenu du cours.",
    ),
]


def register_fonts() -> None:
    font_dir = Path(r"C:\Windows\Fonts")
    fonts = {
        "ClientSans": font_dir / "arial.ttf",
        "ClientSans-Bold": font_dir / "arialbd.ttf",
        "ClientSans-Italic": font_dir / "ariali.ttf",
    }
    for name, path in fonts.items():
        if not path.exists():
            raise FileNotFoundError(f"Police requise introuvable: {path}")
        if name not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont(name, str(path)))


def safe(text: str) -> str:
    return html.escape(text, quote=False)


def paragraph(text: str, style: ParagraphStyle) -> Paragraph:
    return Paragraph(safe(text), style)


def cv_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "name": ParagraphStyle(
            "CvName",
            parent=base["Title"],
            fontName="ClientSans-Bold",
            fontSize=24,
            leading=27,
            textColor=NAVY,
            spaceAfter=2,
        ),
        "title": ParagraphStyle(
            "CvTitle",
            parent=base["Normal"],
            fontName="ClientSans-Bold",
            fontSize=11.5,
            leading=14,
            textColor=TEAL,
            spaceAfter=4,
        ),
        "contact": ParagraphStyle(
            "CvContact",
            parent=base["Normal"],
            fontName="ClientSans",
            fontSize=7.8,
            leading=10.2,
            textColor=TEXT,
            spaceAfter=1,
        ),
        "section": ParagraphStyle(
            "CvSection",
            parent=base["Heading2"],
            fontName="ClientSans-Bold",
            fontSize=11.5,
            leading=13.5,
            textColor=NAVY,
            spaceBefore=7,
            spaceAfter=4,
            borderWidth=0,
        ),
        "body": ParagraphStyle(
            "CvBody",
            parent=base["Normal"],
            fontName="ClientSans",
            fontSize=8.5,
            leading=10.8,
            textColor=TEXT,
            spaceAfter=2,
        ),
        "job": ParagraphStyle(
            "CvJob",
            parent=base["Normal"],
            fontName="ClientSans-Bold",
            fontSize=9.2,
            leading=11,
            textColor=BLUE,
            spaceBefore=4,
            spaceAfter=1,
        ),
        "meta": ParagraphStyle(
            "CvMeta",
            parent=base["Normal"],
            fontName="ClientSans-Italic",
            fontSize=7.7,
            leading=9,
            textColor=MID_GREY,
            spaceAfter=2,
        ),
        "bullet": ParagraphStyle(
            "CvBullet",
            parent=base["Normal"],
            fontName="ClientSans",
            fontSize=7.9,
            leading=9.7,
            textColor=TEXT,
            leftIndent=9,
            firstLineIndent=-7,
            spaceAfter=1.2,
        ),
        "project": ParagraphStyle(
            "CvProject",
            parent=base["Normal"],
            fontName="ClientSans-Bold",
            fontSize=8.8,
            leading=10.5,
            textColor=BLUE,
            spaceBefore=4,
            spaceAfter=1,
        ),
        "small": ParagraphStyle(
            "CvSmall",
            parent=base["Normal"],
            fontName="ClientSans",
            fontSize=7.75,
            leading=9.4,
            textColor=TEXT,
        ),
    }


def cv_footer(canvas, doc) -> None:
    canvas.saveState()
    canvas.setStrokeColor(colors.HexColor("#D0D5DD"))
    canvas.setLineWidth(0.5)
    canvas.line(1.6 * cm, 1.15 * cm, A4[0] - 1.6 * cm, 1.15 * cm)
    canvas.setFont("ClientSans", 7)
    canvas.setFillColor(MID_GREY)
    canvas.drawString(1.6 * cm, 0.72 * cm, "Aymen Khaled - Développeur Full-Stack")
    canvas.drawRightString(A4[0] - 1.6 * cm, 0.72 * cm, f"Page {doc.page}")
    canvas.restoreState()


def cv_header(st: dict[str, ParagraphStyle]) -> list:
    contact_line = (
        f"{PROFILE['location']}  |  {PROFILE['phone']}  |  "
        f"<link href='mailto:{PROFILE['email']}' color='#1D2939'>{PROFILE['email']}</link>"
    )
    links_line = (
        f"<link href='{PROFILE['portfolio']}' color='#1769AA'>{PROFILE['portfolio']}</link>  |  "
        f"<link href='{PROFILE['linkedin']}' color='#1769AA'>LinkedIn</link>  |  "
        f"<link href='{PROFILE['github']}' color='#1769AA'>GitHub</link>"
    )
    return [
        paragraph(PROFILE["name"].upper(), st["name"]),
        paragraph(PROFILE["title"], st["title"]),
        Paragraph(contact_line, st["contact"]),
        Paragraph(links_line, st["contact"]),
        Spacer(1, 3),
    ]


def build_cv_pdf(path: Path) -> None:
    st = cv_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=A4,
        leftMargin=1.6 * cm,
        rightMargin=1.6 * cm,
        topMargin=1.25 * cm,
        bottomMargin=1.45 * cm,
        title="CV Aymen Khaled - Développeur Full-Stack LMS",
        author="Aymen Khaled",
    )
    story: list = []
    story.extend(cv_header(st))
    story.append(paragraph("PROFIL PROFESSIONNEL", st["section"]))
    story.append(paragraph(SUMMARY, st["body"]))
    story.append(paragraph("COMPÉTENCES TECHNIQUES", st["section"]))

    skill_rows = []
    for label, value in SKILLS:
        skill_rows.append(
            [
                Paragraph(f"<b>{safe(label)}</b>", st["small"]),
                paragraph(value, st["small"]),
            ]
        )
    skill_table = Table(skill_rows, colWidths=[3.1 * cm, 14.0 * cm])
    skill_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 0), (0, -1), LIGHT_BLUE),
                ("ROWBACKGROUNDS", (1, 0), (1, -1), [WHITE, LIGHT_GREY]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 2.3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 2.3),
                ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#D0D5DD")),
            ]
        )
    )
    story.append(skill_table)
    story.append(paragraph("EXPÉRIENCE PROFESSIONNELLE", st["section"]))
    for item in EXPERIENCE[:2]:
        block = [
            Paragraph(
                f"{safe(item['role'])} | <font color='#1D2939'>{safe(item['company'])}</font>",
                st["job"],
            ),
            paragraph(f"{item['period']} | {item['location']}", st["meta"]),
        ]
        for bullet in item["bullets"]:
            block.append(Paragraph(f"- {safe(bullet)}", st["bullet"]))
        story.append(KeepTogether(block))

    story.append(PageBreak())
    story.extend(cv_header(st))
    story.append(paragraph("EXPÉRIENCE PROFESSIONNELLE - SUITE", st["section"]))
    for item in EXPERIENCE[2:]:
        block = [
            Paragraph(
                f"{safe(item['role'])} | <font color='#1D2939'>{safe(item['company'])}</font>",
                st["job"],
            ),
            paragraph(f"{item['period']} | {item['location']}", st["meta"]),
        ]
        for bullet in item["bullets"]:
            block.append(Paragraph(f"- {safe(bullet)}", st["bullet"]))
        story.append(KeepTogether(block))
    story.append(paragraph("PROJETS SÉLECTIONNÉS", st["section"]))
    for title, stack, description in PROJECTS:
        story.append(Paragraph(f"{safe(title)} | <font color='#667085'>{safe(stack)}</font>", st["project"]))
        story.append(paragraph(description, st["body"]))

    story.append(paragraph("FORMATION", st["section"]))
    education_rows = []
    for degree, institution, period in EDUCATION:
        education_rows.append(
            [
                Paragraph(f"<b>{safe(degree)}</b><br/><font color='#667085'>{safe(institution)}</font>", st["body"]),
                Paragraph(f"<b>{safe(period)}</b>", st["body"]),
            ]
        )
    education_table = Table(education_rows, colWidths=[14.2 * cm, 2.9 * cm])
    education_table.setStyle(
        TableStyle(
            [
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 0), (-1, -1), [LIGHT_BLUE, WHITE]),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("LINEBELOW", (0, 0), (-1, -1), 0.25, colors.HexColor("#D0D5DD")),
            ]
        )
    )
    story.append(education_table)

    story.append(paragraph("CERTIFICATIONS", st["section"]))
    cert_table = Table(
        [[paragraph(f"- {item}", st["body"])] for item in CERTIFICATIONS],
        colWidths=[17.1 * cm],
    )
    cert_table.setStyle(
        TableStyle(
            [
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(cert_table)

    story.append(paragraph("LANGUES", st["section"]))
    language_table = Table(
        [[paragraph(item, st["body"])] for item in LANGUAGES],
        colWidths=[17.1 * cm],
    )
    language_table.setStyle(
        TableStyle(
            [
                ("LEFTPADDING", (0, 0), (-1, -1), 0),
                ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
            ]
        )
    )
    story.append(language_table)
    story.append(paragraph("Disponibilité", st["section"]))
    story.append(
        paragraph(
            "Disponible pour des missions de développement Full-Stack, d'intégration IA, de modernisation de plateformes SaaS et d'industrialisation de solutions LMS. Fuseau horaire UTC+1.",
            st["body"],
        )
    )
    doc.build(story, onFirstPage=cv_footer, onLaterPages=cv_footer)


def set_cell_text(cell, text: str, bold: bool = False, color: str = "1D2939") -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string(color)


def add_docx_section(document: Document, title: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(11, 46, 79)


def add_docx_bullet(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.45)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
    if not p.runs:
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(8.5)
    else:
        p.runs[0].text = text


def add_docx_experience(document: Document, item: dict) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f"{item['role']} | {item['company']}")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(23, 105, 170)
    meta = document.add_paragraph(f"{item['period']} | {item['location']}")
    meta.paragraph_format.space_after = Pt(1)
    for run in meta.runs:
        run.italic = True
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(102, 112, 133)
    for bullet in item["bullets"]:
        add_docx_bullet(document, bullet)


def build_cv_docx(path: Path) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9)
    normal.paragraph_format.space_after = Pt(2)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(PROFILE["name"].upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 46, 79)

    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(PROFILE["title"])
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(14, 124, 134)

    contact = document.add_paragraph(
        f"{PROFILE['location']} | {PROFILE['phone']} | {PROFILE['email']}\n"
        f"{PROFILE['portfolio']} | {PROFILE['linkedin']} | {PROFILE['github']}"
    )
    contact.paragraph_format.space_after = Pt(4)
    for run in contact.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)

    add_docx_section(document, "Profil professionnel")
    document.add_paragraph(SUMMARY)

    add_docx_section(document, "Compétences techniques")
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in SKILLS:
        row = table.add_row().cells
        set_cell_text(row[0], label, bold=True, color="0B2E4F")
        set_cell_text(row[1], value)
    table.columns[0].width = Cm(3.1)
    table.columns[1].width = Cm(14.0)

    add_docx_section(document, "Expérience professionnelle")
    for item in EXPERIENCE[:2]:
        add_docx_experience(document, item)

    document.add_page_break()
    add_docx_section(document, "Expérience professionnelle - suite")
    for item in EXPERIENCE[2:]:
        add_docx_experience(document, item)
    add_docx_section(document, "Projets sélectionnés")
    for title, stack, description in PROJECTS:
        p = document.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f"{title} | {stack}")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(23, 105, 170)
        document.add_paragraph(description)

    add_docx_section(document, "Formation")
    for degree, institution, period in EDUCATION:
        add_docx_bullet(document, f"{degree} | {institution} | {period}")
    add_docx_section(document, "Certifications")
    for item in CERTIFICATIONS:
        add_docx_bullet(document, item)
    add_docx_section(document, "Langues")
    for item in LANGUAGES:
        add_docx_bullet(document, item)
    add_docx_section(document, "Disponibilité")
    document.add_paragraph(
        "Disponible pour des missions de développement Full-Stack, d'intégration IA, de modernisation de plateformes SaaS et d'industrialisation de solutions LMS. Fuseau horaire UTC+1."
    )
    document.save(path)


def portfolio_styles() -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "cover_kicker": ParagraphStyle(
            "PortfolioCoverKicker",
            parent=base["Normal"],
            fontName="ClientSans-Bold",
            fontSize=13,
            leading=16,
            textColor=colors.HexColor("#BFE5EA"),
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
        "cover_title": ParagraphStyle(
            "PortfolioCoverTitle",
            parent=base["Title"],
            fontName="ClientSans-Bold",
            fontSize=30,
            leading=35,
            textColor=WHITE,
            alignment=TA_CENTER,
            spaceAfter=12,
        ),
        "cover_subtitle": ParagraphStyle(
            "PortfolioCoverSubtitle",
            parent=base["Normal"],
            fontName="ClientSans",
            fontSize=15,
            leading=19,
            textColor=colors.HexColor("#D7EEF1"),
            alignment=TA_CENTER,
            spaceAfter=20,
        ),
        "h1": ParagraphStyle(
            "PortfolioH1",
            parent=base["Heading1"],
            fontName="ClientSans-Bold",
            fontSize=19,
            leading=23,
            textColor=NAVY,
            spaceAfter=10,
        ),
        "h2": ParagraphStyle(
            "PortfolioH2",
            parent=base["Heading2"],
            fontName="ClientSans-Bold",
            fontSize=12,
            leading=15,
            textColor=BLUE,
            spaceBefore=7,
            spaceAfter=4,
        ),
        "body": ParagraphStyle(
            "PortfolioBody",
            parent=base["Normal"],
            fontName="ClientSans",
            fontSize=9.2,
            leading=12.5,
            textColor=TEXT,
            spaceAfter=6,
        ),
        "small": ParagraphStyle(
            "PortfolioSmall",
            parent=base["Normal"],
            fontName="ClientSans",
            fontSize=8,
            leading=10.2,
            textColor=TEXT,
        ),
        "table_header": ParagraphStyle(
            "PortfolioTableHeader",
            parent=base["Normal"],
            fontName="ClientSans-Bold",
            fontSize=8,
            leading=10.2,
            textColor=WHITE,
        ),
        "caption": ParagraphStyle(
            "PortfolioCaption",
            parent=base["Normal"],
            fontName="ClientSans",
            fontSize=9,
            leading=12,
            textColor=TEXT,
            alignment=TA_LEFT,
        ),
    }


def portfolio_page(canvas, doc) -> None:
    canvas.saveState()
    page_w, page_h = landscape(A4)
    if doc.page == 1:
        canvas.setFillColor(NAVY)
        canvas.rect(0, 0, page_w, page_h, fill=1, stroke=0)
        canvas.setFillColor(TEAL)
        canvas.rect(0, 0, page_w, 0.45 * cm, fill=1, stroke=0)
    else:
        canvas.setStrokeColor(colors.HexColor("#D0D5DD"))
        canvas.setLineWidth(0.5)
        canvas.line(1.5 * cm, 1.0 * cm, page_w - 1.5 * cm, 1.0 * cm)
        canvas.setFont("ClientSans", 7.2)
        canvas.setFillColor(MID_GREY)
        canvas.drawString(1.5 * cm, 0.58 * cm, "Portfolio technique LMS - Aymen Khaled")
        canvas.drawRightString(page_w - 1.5 * cm, 0.58 * cm, f"Page {doc.page}")
    canvas.restoreState()


def cropped_screenshot(path: Path) -> io.BytesIO:
    with PILImage.open(path) as image:
        image = image.convert("RGB")
        width, height = image.size
        cropped = image.crop((0, 66, width, height - 38))
        stream = io.BytesIO()
        cropped.save(stream, format="JPEG", quality=92, optimize=True)
        stream.seek(0)
        return stream


def portfolio_table(data: list[list[str]], st: dict[str, ParagraphStyle], widths: list[float]) -> Table:
    rows = []
    for row_index, row in enumerate(data):
        style = st["small"]
        rows.append(
            [
                paragraph(cell, st["table_header"]) if row_index == 0 else paragraph(cell, style)
                for cell in row
            ]
        )
    result = Table(rows, colWidths=widths, repeatRows=1)
    result.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), NAVY),
                ("TEXTCOLOR", (0, 0), (-1, 0), WHITE),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, LIGHT_GREY]),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D0D5DD")),
            ]
        )
    )
    return result


def build_portfolio_pdf(path: Path) -> None:
    st = portfolio_styles()
    page_w, page_h = landscape(A4)
    doc = SimpleDocTemplate(
        str(path),
        pagesize=landscape(A4),
        leftMargin=1.5 * cm,
        rightMargin=1.5 * cm,
        topMargin=1.35 * cm,
        bottomMargin=1.35 * cm,
        title="Portfolio technique LMS - Aymen Khaled",
        author="Aymen Khaled",
    )
    story: list = [
        Spacer(1, 2.2 * cm),
        paragraph("PROJET ATEA", st["cover_kicker"]),
        paragraph("Portfolio technique", st["cover_title"]),
        paragraph("Développement Full-Stack et plateforme LMS", st["cover_subtitle"]),
        Spacer(1, 0.4 * cm),
    ]
    cover_data = [
        ["Profil", "Aymen Khaled - Développeur Full-Stack"],
        ["Spécialités", "Plateformes SaaS, LMS, temps réel, API et intégrations IA"],
        ["Technologies principales", "React, TypeScript, Node.js, FastAPI, PostgreSQL, MongoDB, Redis, Docker"],
        ["Portfolio", PROFILE["portfolio"]],
    ]
    cover_rows = []
    for label, value in cover_data:
        cover_rows.append(
            [
                Paragraph(f"<b><font color='#BFE5EA'>{safe(label)}</font></b>", st["body"]),
                Paragraph(f"<font color='#FFFFFF'>{safe(value)}</font>", st["body"]),
            ]
        )
    cover_table = Table(cover_rows, colWidths=[5.3 * cm, 15.0 * cm])
    cover_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#123B60")),
                ("BOX", (0, 0), (-1, -1), 0.6, colors.HexColor("#2B5F83")),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#2B5F83")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
            ]
        )
    )
    story.append(cover_table)
    story.append(PageBreak())

    story.append(paragraph("1. Positionnement technique", st["h1"]))
    story.append(paragraph(SUMMARY, st["body"]))
    expertise = [
        ["Domaine", "Compétences mobilisables"],
        ["Interfaces", "Espaces apprenant, formateur et administrateur; responsive design; tableaux de bord"],
        ["Services métier", "Cours, leçons, inscriptions, progression, évaluations, certificats et rapports"],
        ["Temps réel", "Classes virtuelles, chat, notifications et synchronisation WebSocket"],
        ["Données", "Modélisation PostgreSQL/MongoDB, cache Redis, recherche, historique et analytique"],
        ["Intégrations", "IA, messagerie, stockage objet, webhooks, OAuth2 et services externes"],
        ["Industrialisation", "Docker, CI/CD, tests, journalisation, supervision, sauvegardes et documentation"],
    ]
    story.append(portfolio_table(expertise, st, [5.1 * cm, 20.6 * cm]))
    story.append(paragraph("Fonctions LMS couvertes", st["h2"]))
    alignment = [
        ["Création", "Diffusion", "Accompagnement", "Évaluation", "Pilotage"],
        [
            "Cours, leçons, médias, plans et ressources",
            "Catalogue, recherche, parcours et accès par rôle",
            "Tutorat, classe virtuelle, chat et notifications",
            "Quiz contextualisés, suivi des résultats et progression",
            "Administration, utilisateurs, indicateurs et rapports",
        ],
    ]
    story.append(portfolio_table(alignment, st, [5.14 * cm] * 5))
    story.append(PageBreak())

    story.append(paragraph("2. Réalisations sélectionnées", st["h1"]))
    references = [
        ["Réalisation", "Rôle et technologies", "Périmètre"],
        [
            "EduNova - plateforme e-learning",
            "Développeur Full-Stack | React, Node.js, MongoDB, WebSockets, IA",
            "Cours et leçons, studio de production, sessions synchrones, quiz, catalogue, espaces apprenant et administration.",
        ],
        [
            "Strategy Navigator",
            "Développeur Web Full-Stack | React, Node.js, Redis, RBAC, API",
            "SaaS multi-tenant, isolation des données, tableaux de bord, 65+ services IA, paiements et intégrations collaboratives.",
        ],
        [
            "JourneyAI",
            "Développeur Web Full-Stack | WebSockets, Whisper, OpenAI, Groq, Stripe",
            "Plateforme de plus de dix assistants IA, conversation temps réel, transcription et facturation à l'usage.",
        ],
        [
            "SaleSide AI",
            "Développeur Web Full-Stack | IA conversationnelle, transcription, API",
            "Assistant de réunion commerciale: diarisation, suivi des objections, synthèses et actions recommandées.",
        ],
        [
            "RxCare / Deep Analyzer Suite",
            "Développeur Full-Stack mobile | React Native, Expo, Express, OpenAI",
            "Suivi des prescriptions, interactions médicamenteuses, téléconsultation, messagerie sécurisée et assistant IA.",
        ],
        [
            "Integrity Solve",
            "Développeur Full-Stack | React, TypeScript, Express, PostgreSQL",
            "Onboarding, dossiers de conformité AML/CTF, validations, escalades et synthèses d'audit PDF.",
        ],
    ]
    story.append(portfolio_table(references, st, [5.2 * cm, 8.4 * cm, 12.1 * cm]))
    story.append(paragraph("Liens publics", st["h2"]))
    links = [
        ["Portfolio", PROFILE["portfolio"]],
        ["EduNova", "https://e-learning-five-tau.vercel.app"],
        ["Strategy Navigator", "https://strategynavigator.ai"],
        ["JourneyAI", "https://app.meetjourney.ai"],
        ["SaleSide AI", "https://saleside.ai"],
        ["GitHub", PROFILE["github"]],
    ]
    story.append(portfolio_table([["Produit", "Adresse"]] + links, st, [5.2 * cm, 20.5 * cm]))
    story.append(PageBreak())

    story.append(paragraph("3. Architecture technique proposée", st["h1"]))
    story.append(
        paragraph(
            "L'architecture s'appuie sur une base Full-Stack moderne et modulaire. Les contenus sont administrables depuis l'interface; les équipes pédagogiques peuvent structurer les parcours, importer les ressources et configurer les évaluations sans intervenir dans le code.",
            st["body"],
        )
    )
    architecture = [
        ["Couche", "Technologies", "Responsabilités"],
        ["Interfaces web", "React, TypeScript, Next.js", "Espaces par rôle, accessibilité, responsive design et tableaux de bord"],
        ["API métier", "Node.js, Express ou FastAPI", "Utilisateurs, cours, progression, évaluations, certificats, rapports et intégrations"],
        ["Données", "PostgreSQL, MongoDB, Redis", "Données métier, contenus flexibles, cache, sessions et analytique"],
        ["Médias", "Stockage objet compatible S3", "Vidéos, PDF, guides, exercices, pièces jointes et productions apprenants"],
        ["Temps réel", "WebSockets", "Classes virtuelles, chat, présence, notifications et collaboration"],
        ["Sécurité", "TLS, JWT/OAuth2, RBAC", "Authentification, autorisations, validation, journalisation et traçabilité"],
        ["Exploitation", "Docker, CI/CD, Nginx", "Staging, production, tests, supervision, sauvegardes et mises à jour"],
    ]
    story.append(portfolio_table(architecture, st, [4.0 * cm, 6.0 * cm, 15.7 * cm]))
    story.append(paragraph("Cycle de mise en œuvre", st["h2"]))
    phases = [
        ["1. Cadrage", "2. Audit", "3. Adaptation", "4. Validation", "5. Déploiement"],
        [
            "Besoins, rôles, parcours et critères d'acceptation",
            "Code, dépendances, données, sécurité et performance",
            "Fonctions, interfaces, intégrations et migration",
            "Tests fonctionnels, sécurité, charge et recette",
            "Production, supervision, documentation et transfert",
        ],
    ]
    story.append(portfolio_table(phases, st, [5.14 * cm] * 5))
    story.append(PageBreak())

    image_streams: list[io.BytesIO] = []
    for index, (filename, title, caption) in enumerate(SCREENSHOTS, 1):
        source = SCREENSHOT_DIR / filename
        if not source.exists():
            raise FileNotFoundError(f"Capture LMS introuvable: {source}")
        story.append(paragraph(f"4.{index} {title}", st["h1"]))
        stream = cropped_screenshot(source)
        image_streams.append(stream)
        with PILImage.open(stream) as screenshot:
            iw, ih = screenshot.size
        stream.seek(0)
        max_w = page_w - 3.0 * cm
        max_h = page_h - 6.0 * cm
        scale = min(max_w / iw, max_h / ih)
        screenshot_flowable = Image(stream, width=iw * scale, height=ih * scale)
        screenshot_flowable.hAlign = "CENTER"
        story.append(screenshot_flowable)
        story.append(Spacer(1, 0.25 * cm))
        caption_box = Table(
            [
                [
                    Paragraph("<b>Fonction présentée</b>", st["caption"]),
                    paragraph(caption, st["caption"]),
                ]
            ],
            colWidths=[4.0 * cm, 21.7 * cm],
        )
        caption_box.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), LIGHT_BLUE),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 8),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                    ("TOPPADDING", (0, 0), (-1, -1), 7),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
                    ("BOX", (0, 0), (-1, -1), 0.4, colors.HexColor("#B8D5E5")),
                ]
            )
        )
        story.append(caption_box)
        if index < len(SCREENSHOTS):
            story.append(PageBreak())

    doc.build(story, onFirstPage=portfolio_page, onLaterPages=portfolio_page)


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def prepare_output() -> None:
    if CLIENT_DIR.exists():
        resolved_client = CLIENT_DIR.resolve()
        resolved_output = OUTPUT_ROOT.resolve()
        if resolved_output not in resolved_client.parents:
            raise RuntimeError(f"Refus de supprimer un chemin hors sortie: {resolved_client}")
        shutil.rmtree(resolved_client)
    CLIENT_DIR.mkdir(parents=True, exist_ok=True)
    CLIENT_DIPLOMA.parent.mkdir(parents=True, exist_ok=True)
    EDITABLE_DIR.mkdir(parents=True, exist_ok=True)
    if CLIENT_ZIP.exists():
        CLIENT_ZIP.unlink()


def copy_diploma() -> None:
    if not SOURCE_DIPLOMA.exists():
        raise FileNotFoundError(f"Diplôme introuvable: {SOURCE_DIPLOMA}")
    shutil.copy2(SOURCE_DIPLOMA, CLIENT_DIPLOMA)
    if sha256(SOURCE_DIPLOMA) != sha256(CLIENT_DIPLOMA):
        raise RuntimeError("La copie du diplôme n'est pas identique au fichier source")


def build_zip() -> None:
    files = [CLIENT_CV, CLIENT_PORTFOLIO, CLIENT_DIPLOMA]
    with zipfile.ZipFile(CLIENT_ZIP, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for source in files:
            archive.write(source, arcname=source.relative_to(CLIENT_DIR).as_posix())


def main() -> None:
    register_fonts()
    prepare_output()
    build_cv_pdf(CLIENT_CV)
    build_cv_docx(EDITABLE_CV)
    build_portfolio_pdf(CLIENT_PORTFOLIO)
    copy_diploma()
    build_zip()
    print(f"Dossier client généré: {CLIENT_ZIP}")


if __name__ == "__main__":
    main()
